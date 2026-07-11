from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mohammedraad/Desktop/AiQo")
OUT = ROOT / "output" / "AiQo_التوثيق_الشامل_2026-07-03.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "152238"
INK = "1F2933"
MUTED = "5B6775"
PALE_BLUE = "E8EEF5"
PALE_MINT = "E8F7F2"
MINT = "5ECDB7"
PALE_GOLD = "FFF5DF"
GOLD = "D9A441"
PALE_RED = "FDECEC"
RED = "B54747"
WHITE = "FFFFFF"
LIGHT_GREY = "F5F7FA"
MID_GREY = "D8DEE7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color: str = MID_GREY, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    bidi_visual = tbl_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi_visual)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph_rtl(paragraph, rtl: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl:
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            p_pr.append(bidi)
        bidi.set(qn("w:val"), "1")
    elif bidi is not None:
        p_pr.remove(bidi)


def run_rtl(run, rtl: bool = True) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl_node = r_pr.find(qn("w:rtl"))
    if rtl:
        if rtl_node is None:
            rtl_node = OxmlElement("w:rtl")
            r_pr.append(rtl_node)
        rtl_node.set(qn("w:val"), "1")
    elif rtl_node is not None:
        r_pr.remove(rtl_node)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), "Arial")


def style_font(style, name: str, size: float, color: str = INK, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), name)


def set_style_bidi(style) -> None:
    p_pr = style.element.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    style_font(normal, "Arial", 11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True
    set_style_bidi(normal)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style_font(style, "Arial", size, color, True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        set_style_bidi(style)

    title = styles["Title"]
    style_font(title, "Arial", 28, NAVY, True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    set_style_bidi(title)

    subtitle = styles.add_style("AiQo Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    style_font(subtitle, "Arial", 15, BLUE, False)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtitle.paragraph_format.space_after = Pt(18)
    set_style_bidi(subtitle)

    deck = styles.add_style("AiQo Deck", WD_STYLE_TYPE.PARAGRAPH)
    style_font(deck, "Arial", 11, MUTED, False)
    deck.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    deck.paragraph_format.space_after = Pt(8)
    deck.paragraph_format.line_spacing = 1.25
    set_style_bidi(deck)

    part = styles.add_style("AiQo Part", WD_STYLE_TYPE.PARAGRAPH)
    style_font(part, "Arial", 21, NAVY, True)
    part.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    part.paragraph_format.space_before = Pt(28)
    part.paragraph_format.space_after = Pt(12)
    part.paragraph_format.keep_with_next = True
    set_style_bidi(part)

    callout = styles.add_style("AiQo Callout", WD_STYLE_TYPE.PARAGRAPH)
    style_font(callout, "Arial", 10.5, DARK_BLUE, False)
    callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.left_indent = Inches(0.16)
    callout.paragraph_format.right_indent = Inches(0.16)
    callout.paragraph_format.line_spacing = 1.2
    set_style_bidi(callout)

    flow = styles.add_style("AiQo Flow", WD_STYLE_TYPE.PARAGRAPH)
    style_font(flow, "Arial", 10.5, NAVY, True)
    flow.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow.paragraph_format.space_before = Pt(8)
    flow.paragraph_format.space_after = Pt(10)
    flow.paragraph_format.line_spacing = 1.25
    set_style_bidi(flow)

    small = styles.add_style("AiQo Small", WD_STYLE_TYPE.PARAGRAPH)
    style_font(small, "Arial", 8.5, MUTED, False)
    small.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.1
    set_style_bidi(small)


def add_bottom_border(paragraph, color: str = BLUE, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep])
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    run._r.append(fld_char_end)


def add_run(paragraph, text: str, *, bold: bool = False, color: str | None = None, size: float | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    run_rtl(run)
    return run


def p(doc: Document, text: str = "", style: str | None = None, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style in (None, "Normal") else paragraph.alignment
    paragraph_rtl(paragraph)
    if bold_lead and text.startswith(bold_lead):
        add_run(paragraph, bold_lead, bold=True)
        add_run(paragraph, text[len(bold_lead):])
    else:
        add_run(paragraph, text)
    return paragraph


def paras(doc: Document, text: str) -> None:
    for block in [b.strip() for b in text.strip().split("\n\n") if b.strip()]:
        p(doc, " ".join(line.strip() for line in block.splitlines()))


def heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    paragraph_rtl(paragraph)
    for run in paragraph.runs:
        run_rtl(run)
    return paragraph


def part(doc: Document, number: str, title: str, strapline: str) -> None:
    marker_text = "مرجع سريع" if number == "—" else f"القسم {number}"
    marker = p(doc, marker_text, "AiQo Small")
    marker.paragraph_format.space_before = Pt(40)
    marker.paragraph_format.keep_with_next = True
    marker.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    marker.runs[0].bold = True
    title_p = p(doc, title, "AiQo Part")
    add_bottom_border(title_p, MINT, 20)
    strap = p(doc, strapline, "AiQo Deck")
    strap.paragraph_format.keep_with_next = True


def callout(doc: Document, title: str, text: str, tone: str = "blue"):
    fill, color = {
        "blue": (PALE_BLUE, DARK_BLUE),
        "mint": (PALE_MINT, "216A5A"),
        "gold": (PALE_GOLD, "7A5916"),
        "red": (PALE_RED, RED),
    }[tone]
    paragraph = doc.add_paragraph(style="AiQo Callout")
    paragraph_rtl(paragraph)
    shade_paragraph(paragraph, fill)
    add_run(paragraph, f"{title}: ", bold=True, color=color)
    add_run(paragraph, text, color=color)
    return paragraph


def flow(doc: Document, text: str) -> None:
    paragraph = p(doc, text, "AiQo Flow")
    shade_paragraph(paragraph, PALE_BLUE)
    add_bottom_border(paragraph, MINT, 10)


def add_numbering_definitions(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element

    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def abstract_num(abstract_id: int, fmt: str, text: str, font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "right")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:right"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def num_instance(num_id: int, abstract_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    abstract_num(next_abs, "bullet", "•", "Arial")
    num_instance(next_num, next_abs)
    bullet_num = next_num

    abstract_num(next_abs + 1, "decimal", "%1.")
    num_instance(next_num + 1, next_abs + 1)
    decimal_num = next_num + 1
    return bullet_num, decimal_num


def list_item(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph_rtl(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)
    add_run(paragraph, text)


def bullets(doc: Document, items: Iterable[str], bullet_num: int) -> None:
    for item in items:
        list_item(doc, item, bullet_num)


def numbered(doc: Document, items: Iterable[str], decimal_num: int) -> None:
    for item in items:
        list_item(doc, item, decimal_num)


def table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    header_fill: str = PALE_BLUE,
    stripe: bool = True,
    font_size: float = 9.0,
):
    tbl = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(tbl)
    hdr = tbl.rows[0]
    set_repeat_table_header(hdr)
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        para = cell.paragraphs[0]
        paragraph_rtl(para)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_after = Pt(0)
        run = add_run(para, value, bold=True, color=NAVY, size=font_size)
        run.font.name = "Arial"

    for row_idx, values in enumerate(rows):
        row = tbl.add_row()
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            if stripe and row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY)
            para = cell.paragraphs[0]
            paragraph_rtl(para)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.1
            add_run(para, str(value), size=font_size)
    set_table_geometry(tbl, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return tbl


def setup_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_styles(doc)
    bullet_num, decimal_num = add_numbering_definitions(doc)

    header = section.header
    hp = header.paragraphs[0]
    paragraph_rtl(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(hp, "AiQo  •  الدليل الشامل والعميق", bold=True, color=DARK_BLUE, size=8.5)
    add_bottom_border(hp, MINT, 8)

    footer = section.footer
    fp = footer.paragraphs[0]
    paragraph_rtl(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, "وثيقة مرجعية  •  3 يوليو 2026  •  صفحة ", color=MUTED, size=8)
    add_field(fp, "PAGE", "1")

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    paragraph_rtl(ffp)
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(ffp, "AiQo • توثيق المنتج والمنظومة التقنية", color=MUTED, size=8)

    props = doc.core_properties
    props.title = "AiQo — الدليل الشامل والعميق للتطبيق والمنظومة التقنية"
    props.subject = "توثيق المنتج، تجربة المستخدم، الذكاء الاصطناعي، البيانات، البنية، الأمان، والواقع التشغيلي"
    props.author = "OpenAI Codex — مبني على فحص مستودع AiQo"
    props.keywords = "AiQo, iOS, watchOS, HealthKit, Apple Intelligence, Captain, Kernel, Gym, Kitchen"
    props.comments = "لقطة موثقة بتاريخ 2026-07-03، مع فصل الواقع المنفذ عن الرؤية التسويقية."
    return doc, bullet_num, decimal_num


def add_cover(doc: Document) -> None:
    accent = p(doc, "AiQo", "AiQo Small")
    accent.paragraph_format.space_before = Pt(34)
    accent.runs[0].font.size = Pt(15)
    accent.runs[0].font.bold = True
    accent.runs[0].font.color.rgb = RGBColor.from_string(MINT)
    add_bottom_border(accent, MINT, 28)

    title = doc.add_paragraph(style="Title")
    paragraph_rtl(title)
    add_run(title, "الدليل الشامل والعميق\nلتطبيق AiQo", bold=True, color=NAVY, size=28)

    sub = p(
        doc,
        "توثيق المنتج، تجربة المستخدم، الذكاء الاصطناعي، البيانات الصحية، البنية التقنية، الأمان، الاشتراكات، والواقع التشغيلي",
        "AiQo Subtitle",
    )
    sub.paragraph_format.space_after = Pt(26)

    p(doc, "لقطة مرجعية من المستودع الحالي", "AiQo Deck")
    meta = [
        "تاريخ التحقق: 3 يوليو 2026",
        "نسخة المشروع: 1.0.8 — Build 31",
        "النسخة الظاهرة في متجر الإمارات وقت التحقق: 1.0.7",
        "الفرع المفحوص: program/world-class-completion",
        "النطاق: iPhone + Apple Watch + Widgets + Live Activities + Kernel + Web + Supabase",
    ]
    for line in meta:
        paragraph = p(doc, line, "AiQo Deck")
        paragraph.paragraph_format.space_after = Pt(5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    callout(
        doc,
        "قاعدة هذه الوثيقة",
        "كل ادعاء مهم صُنّف وفق ما يثبته الكود والإعدادات الحالية. عندما يختلف التنفيذ عن الموقع أو وثائق أقدم، تكون الأولوية للحقيقة التنفيذية، مع تسجيل الفرق صراحة.",
        "mint",
    )
    p(
        doc,
        "إعداد تحليلي داخلي • للاستخدام في المنتج والهندسة والتصميم والخصوصية والاستثمار والدعم",
        "AiQo Small",
    )


def add_front_matter(doc: Document, bullet_num: int) -> None:
    doc.add_page_break()
    heading(doc, "كيف تقرأ هذه الوثيقة", 1)
    paras(
        doc,
        """
        هذه ليست إعادة صياغة لصفحة المتجر، وليست مواصفات مستقبلية. هي خريطة موحّدة لما يفعله AiQo الآن، وكيف ينتقل المستخدم خلاله، وأين تُخزن بياناته، وما الذي يعمل محليًا أو سحابيًا، وما الذي تحكمه الاشتراكات والأعلام التجريبية، وأين توجد فجوات أو ازدواجيات تحتاج قرارًا.

        شمل الفحص كود Swift وSwiftUI، إعدادات Xcode والـ entitlements، ملفات الخصوصية، StoreKit، Edge Functions، تطبيق Apple Watch، الويدجت والامتدادات، موقع aiqo.app، ووثائق المشروع. الأرقام الخاصة بالمتجر والأسعار موثقة من واجهة متجر الإمارات بتاريخ الوثيقة، وقد تتغير لاحقًا.
        """,
    )
    heading(doc, "مفاتيح الحالة", 2)
    bullets(
        doc,
        [
            "منفّذ ومفعّل: يوجد مسار قابل للتشغيل في النسخة الحالية ولا يمنعه علم افتراضي مغلق.",
            "منفّذ خلف علم: الكود موجود، لكن ظهوره أو مساره يعتمد على Feature Flag أو إعداد Build.",
            "منفّذ وغير ظاهر: بنية أو واجهة موجودة لكن الوصول العام إليها مغلق، مثل Tribe في الإعداد الحالي.",
            "اعتماد تشغيلي خارجي: يحتاج إعدادًا أو خدمة خارج المستودع، مثل أسرار Supabase أو مزود الصوت.",
            "رؤية أو وثيقة أقدم: وصف موجود في الموقع أو ملفات Markdown لكنه لا يطابق التنفيذ الحالي بالكامل.",
        ],
        bullet_num,
    )
    callout(
        doc,
        "تنبيه طبي",
        "AiQo تطبيق صحة ولياقة وإرشاد عام، وليس جهازًا طبيًا ولا بديلًا عن الطبيب أو خدمات الطوارئ. الاستنتاجات الصحية داخل الكابتن والدرجات اليومية توجيهية، ويجب أن تبقى لغة المنتج ضمن هذا الحد.",
        "gold",
    )

    heading(doc, "فهرس المحتويات", 1)
    toc = doc.add_paragraph()
    paragraph_rtl(toc)
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "يُحدّث الفهرس عند فتح الملف في Microsoft Word")
    p(
        doc,
        "إن لم يظهر الفهرس تلقائيًا: حدده داخل Word ثم اختر «تحديث الحقل». العناوين نفسها مرتبة أيضًا عبر الأقسام الاثني عشر والملاحق.",
        "AiQo Small",
    )


def _add_executive_snapshot_base(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "الأول",
        "ما هو AiQo؟",
        "تعريف المنتج، الوعد الأساسي، بنية التجربة، واللقطة التنفيذية التي يجب أن يعرفها أي شخص قبل الغوص في التفاصيل.",
    )
    heading(doc, "الخلاصة التنفيذية", 1)
    paras(
        doc,
        """
        AiQo هو نظام صحة ولياقة شخصي عربي أولًا، مبني حول فكرة «نظام تشغيل حيوي رقمي»: يجمع إشارات الحركة والنوم والترطيب والتمارين والمزاج، ثم يحوّلها إلى واجهات يومية، تحديات، وخطط، وإرشاد حواري من شخصية «الكابتن». التطبيق لا يقتصر على تسجيل البيانات؛ طموحه أن يربط الإدراك بالفعل: يرى المستخدم وضعه، يفهم السبب، ثم ينفذ خطوة قابلة للقياس.

        القلب الحالي للمنتج أربع تبويبات رئيسية: الرئيسية، النادي، المطبخ، والكابتن. حولها توجد طبقات قوية لكنها أقل ظهورًا: النوم، الترطيب الذكي، My Vibe، Kernel للتحكم في المشتتات، Legendary Peaks، الصور التقدمية، التقارير الأسبوعية، Apple Watch، الويدجت، وLive Activities. توجد كذلك منظومة ويب عامة وطبقة Supabase، بينما Tribe مبني بدرجة كبيرة لكنه مخفي افتراضيًا.

        الميزة الفارقة تقنيًا هي وجود مسارين للكابتن. المستوى المجاني يستخدم Apple Foundation Models على الجهاز عندما تتوفر Apple Intelligence؛ لا يرسل المحادثة إلى السحابة، ويحافظ على سياق الجلسة داخل LanguageModelSession. المستويات المدفوعة تستخدم عقلًا أوسع: سحابة Gemini، ذاكرة دائمة متعددة الأنواع، استرجاع، توجيهات قابلة للتنفيذ، تخصيص، بطاقات منظمة، وصوت مميز. لذلك التقسيم الحقيقي ليس «دردشة أو لا دردشة»، بل «محادثة محلية سريعة» مقابل «علاقة مستمرة عميقة».
        """,
    )

    table(
        doc,
        ["البعد", "الحقيقة الحالية"],
        [
            ["المنصة الأساسية", "iOS 26.2+ مع تطبيق watchOS 26.2+ وامتدادات وويدجت."],
            ["اللغات", "العربية افتراضيًا والإنجليزية؛ 2730+ مفتاحًا لكل لغة تقريبًا."],
            ["التنقل الرئيسي", "4 تبويبات: الرئيسية، النادي، المطبخ، الكابتن."],
            ["الذكاء المجاني", "Apple Intelligence محلي، سياق جلسة، بيانات صحية حية، بلا ذاكرة دائمة."],
            ["الذكاء المدفوع", "Gemini + ذاكرة واسترجاع وتخصيص وتوجيهات وصوت وبطاقات منظمة."],
            ["البيانات الصحية", "HealthKit قراءة/كتابة مع ملخصات يومية وتاريخية ومراقبة خلفية."],
            ["الاشتراكات", "Free، Max، Pro؛ والتجربة StoreKit تُعامل وظيفيًا كـ Pro."],
            ["المتجر وقت التحقق", "الإصدار 1.0.7؛ Max بسعر 39.99 د.إ وPro بسعر 79.99 د.إ شهريًا."],
            ["المشروع الحالي", "1.0.8 Build 31، يتضمن توسيع المجاني والكابتن المحلي وKernel المجاني لتطبيق واحد."],
            ["الحالة الطبية", "إرشاد صحة ولياقة، وليس تشخيصًا أو جهازًا طبيًا."],
        ],
        [2340, 7020],
    )

    heading(doc, "الوعد المنتجّي", 2)
    bullets(
        doc,
        [
            "تجميع الحياة الصحية في مكان واحد بدل تطبيق لكل مقياس.",
            "تحويل البيانات الخام إلى معنى يومي مفهوم باللغة والثقافة الأقرب للمستخدم.",
            "تقليل الاحتكاك: حلقات، اختصارات، Watch، ويدجت، Live Activities، وإجراءات من الإشعار.",
            "صناعة استمرارية طويلة عبر مستويات XP، سلاسل، تحديات، خطط، وذاكرة الكابتن.",
            "الخصوصية المتدرجة: تنفيذ محلي عندما يكفي، وموافقة واضحة قبل أي مسار سحابي.",
            "ربط الصحة بالانتباه الرقمي عبر Kernel، وبالمزاج والصوت عبر My Vibe.",
        ],
        bullet_num,
    )


def add_user_journey(doc: Document, bullet_num: int, decimal_num: int) -> None:
    part(
        doc,
        "الثاني",
        "رحلة المستخدم وتجربة التطبيق",
        "من أول تشغيل إلى الاستخدام اليومي: التسجيل، الموافقات، الشاشة الرئيسية، التنقل، والعودة من الإشعارات والاختصارات.",
    )
    heading(doc, "المسار الكامل لأول تشغيل", 1)
    flow(
        doc,
        "اللغة ← تسجيل الدخول أو الضيف ← الملف الشخصي ← HealthKit التاريخي ← موافقة AI ← إخلاء طبي ← الإعداد السريع ← تعريف الميزات ← عرض الاشتراك ← التطبيق",
    )
    numbered(
        doc,
        [
            "اختيار اللغة: العربية هي الافتراضية، والإنجليزية متاحة. يضبط AppSettingsStore لغة الحزمة واتجاه الواجهة.",
            "الدخول: Sign in with Apple عبر Supabase OIDC وnonce، أو الاستمرار كضيف. الضيف يحصل على تجربة محلية لكنه لا يملك مزايا الحساب المتزامن نفسها.",
            "الملف الأساسي: الاسم والعمر والوزن والطول والجنس والهدف. بعض البيانات تستخدم لتخصيص الواجهة والحسابات، وبعضها يصبح سياقًا للكابتن فقط بعد الموافقة.",
            "البيانات التاريخية: طلب HealthKit اختياري، مع إمكانية التخطي. الغرض ملء الرئيسية والتقارير بدل البدء من يوم فارغ.",
            "موافقة الذكاء: تشرح انتقال بعض المدخلات إلى مزود سحابي في المسارات المدفوعة. يمكن رفضها والعمل بوضع محلي فقط، ويمكن سحبها لاحقًا.",
            "إخلاء المسؤولية الطبية: يثبت أن التطبيق إرشادي ولا يعالج الطوارئ أو يشخص المرض.",
            "الإعداد السريع: الهدف، الرياضة، وقت التمرين، سنة الميلاد، إشارات السلامة الصحية، وأوقات النوم والاستيقاظ.",
            "التحقق العمري: الحد الأدنى 18 سنة. المسار يمنع إكمال الإعداد لمن هم دون ذلك.",
            "تعريف الميزات: صفحتان قصيرتان لتقديم القيمة الأساسية.",
            "عرض الاشتراك: يظهر قبل الرئيسية لكنه قابل للتخطي في التنفيذ الحالي؛ لذلك لا يوجد Hard Paywall فعلي عند الإطلاق.",
        ],
        decimal_num,
    )
    callout(
        doc,
        "ملاحظة هجرة",
        "المستخدمون القدامى لا يُجبرون آليًا على إعادة كل الشاشات. AppFlowController يحمل منطقًا لتجاوز مراحل الإعداد الجديدة أو تعليمها مكتملة، حتى لا ينكسر مسار من استخدم إصدارات سابقة.",
        "blue",
    )

    heading(doc, "الهوية والحساب", 2)
    paras(
        doc,
        """
        الحساب السحابي مبني فوق Supabase Auth. تسجيل Apple يستخدم nonce لتقليل إعادة تشغيل الرمز، ثم تُربط هوية المستخدم بملف profiles وبخدمات تقارير الأعطال عند توفر Firebase. تسجيل الخروج يمسح ارتباط التحليلات والهوية المحلية الملائمة، بينما حذف الحساب موجود في الإعدادات كمتطلب خصوصية وتشغيلي.

        وضع الضيف مهم لتقليل الاحتكاك، لكنه يعني أن الاستمرارية تعتمد أكثر على UserDefaults وSwiftData وHealthKit على الجهاز. لا ينبغي تسويق الضيف كنسخة احتياطية سحابية؛ هو مسار استخدام محلي دون ضمان استرجاع البيانات بعد حذف التطبيق أو تبديل الجهاز.
        """,
    )

    heading(doc, "التنقل الرئيسي", 1)
    table(
        doc,
        ["التبويب", "الدور", "أهم المداخل"],
        [
            ["الرئيسية", "لوحة اليوم والاتجاهات.", "الحلقات، الهالة، الماء، السلسلة، My Vibe، الملف."],
            ["النادي", "مركز الحركة والتدريب.", "Body، Plan، Peaks، Battle، Impact."],
            ["المطبخ", "التغذية والمخزون والخطط.", "وجبات اليوم، الماكروز، الثلاجة، الخطة، التسوق."],
            ["الكابتن", "الحوار والتحليل والذاكرة.", "رسالة، صورة، صوت، بطاقات، تخصيص، ملاحظات وتذكيرات."],
        ],
        [1560, 3120, 4680],
    )
    paras(
        doc,
        """
        MainTabRouter يعرّف أربعة تبويبات فعلية. هذه نقطة مهمة لأن بعض وثائق المشروع القديمة ما زالت تصف ثلاثة تبويبات أو تضع الكابتن كزر عائم. الواقع الحالي: الكابتن تبويب مستقل، والمطبخ كذلك. توجد شاشات عميقة تُفتح من الملف أو الروابط العميقة دون أن تكون تبويبات، مثل النوم وKernel والتقرير الأسبوعي.
        """,
    )

    heading(doc, "العودة إلى التطبيق", 2)
    bullets(
        doc,
        [
            "الإشعارات المحلية تفتح سياقًا محددًا أو تمرر رسالة إلى الكابتن عبر CaptainNotificationHandler.",
            "الروابط العميقة تستخدم aiqo://، ويوجد مسار aiqo-spotify لمصادقة Spotify.",
            "Siri وNSUserActivity يدعمان بدء المشي والجري وHIIT، فتح الكابتن، ملخص اليوم، تسجيل الماء، المطبخ، والتقرير الأسبوعي.",
            "Live Activity تعيد المستخدم إلى جلسة التمرين الحية، بينما الويدجت يعرض حالة مختصرة أو يسجل 250 مل ماء.",
            "Apple Watch يرسل لقطات الجلسة والملخصات عبر WatchConnectivity، ويمكن للهاتف إكمال العرض والتحليل.",
        ],
        bullet_num,
    )

    heading(doc, "مبادئ تجربة الاستخدام", 2)
    bullets(
        doc,
        [
            "قيمة في أول نظرة: الحلقات والملخص قبل الجداول المعقدة.",
            "إجراء واحد واضح: تسجيل ماء، بدء تمرين، أو سؤال الكابتن.",
            "العمق اختياري: الخطة والذاكرة والتقارير لا تعيق المستخدم البسيط.",
            "اللغة محلية ثقافيًا، لا ترجمة حرفية لمنتج إنجليزي.",
            "تقليل الإحراج: بيانات الصحة والصور والمزاج وراء موافقات وسياقات واضحة.",
            "الاستمرارية قابلة للرؤية: XP، المستوى، السلسلة، الإنجازات، والمقارنة.",
        ],
        bullet_num,
    )

    heading(doc, "الاتجاه البصري وإمكانية الوصول", 2)
    paras(
        doc,
        """
        نظام التصميم يستخدم مسافات 8 و12 و16 و24 نقطة، وأنصاف أقطار 12 و16 و24، وحدًا أدنى للمس 44 نقطة. الخلفيات فاتحة وداكنة، مع نعناع #5ECDB7 كلون إبراز، ودرجات زجاجية وبيج وذهبي للمكافآت. الخطوط Rounded من منظومة Apple، مع عناصر Glass عندما يسمح النظام.

        توجد أدوات لـ Dynamic Type، VoiceOver، Reduce Motion، وReduce Transparency. لكن صفحة المتجر وقت التحقق لا تعلن ميزات وصول رسميًا؛ أي أن وجود دعم في الكود لا يساوي اكتمال التدقيق أو التصريح. يلزم اختبار قارئ الشاشة والتكبير والتباين في كل مسار مدفوع وحساس.
        """,
    )


def add_home_health_sleep(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "الثالث",
        "الرئيسية، HealthKit، الماء، والنوم",
        "الطبقة التي تحول القياسات الصحية إلى صورة يومية قابلة للفهم والفعل.",
    )
    heading(doc, "الرئيسية: مركز اليوم", 1)
    paras(
        doc,
        """
        HomeView ليست قائمة أرقام؛ هي لوحة قيادة يومية. تجمع الخطوات، السعرات النشطة، نسبة الوقوف، الماء، النوم، والمسافة، وتعرضها كحلقات وبطاقات ورسوم حسب النطاق: يوم، أسبوع، شهر، سنة، وكل الوقت. HomeViewModel يقرأ HealthKit ويحدّث الواجهة، بينما snapshot مختصر يذهب إلى App Group للويدجت.

        «Daily Aura» طبقة شعورية بصرية فوق الأرقام: هدف وهالة وتاريخ بسيط. بجانبها يظهر زجاجة أو حلقة الماء، شارة السلسلة، احتفال المستوى عند الارتقاء، مدخل الملف، وبطاقة My Vibe. هذا الترتيب يجعل الرئيسية جسرًا بين القياس، التحفيز، والحالة النفسية.
        """,
    )
    table(
        doc,
        ["المؤشر", "المصدر", "طريقة العرض أو الاستخدام"],
        [
            ["الخطوات", "HealthKit stepCount", "حلقة، اتجاهات، XP/عملات، Kernel، مهام."],
            ["السعرات النشطة", "activeEnergyBurned", "حلقة، ملخص، حساب XP، تعافي."],
            ["الوقوف", "standHour / Activity Summary", "نسبة إنجاز يومية."],
            ["الماء", "dietaryWater + إدخال التطبيق", "مل، لتر، حلقة، وتيرة وتذكير."],
            ["النوم", "sleepAnalysis", "ساعات، مراحل، درجة، اتجاهات، Smart Wake."],
            ["المسافة", "walkingRunningDistance وغيرها", "كم يومي واتجاهات وجري."],
            ["القلب والتعافي", "HR، Resting HR، HRV، VO₂ max", "سياق للكابتن وخطط وتعافي."],
        ],
        [1870, 2800, 4690],
    )

    heading(doc, "طبقة HealthKit", 1)
    paras(
        doc,
        """
        HealthKitService actor هو الواجهة الحديثة الأساسية لتقليل سباقات القراءة والكتابة. يطلب أقل قدر يلزم لكل مرحلة، ويستخدم anchored/background queries لبعض التحديثات. توجد أيضًا HealthKitManager أقدم ما زالت له استدعاءات، وهو ازدواج تقني ينبغي توحيده حتى لا تختلف الصلاحيات أو طرق التجميع.

        الأنواع المقروءة تشمل الخطوات، نبض القلب، النبض في الراحة، HRV، متوسط نبض المشي، الطاقة النشطة، مسافات المشي والجري والدراجات، ماء الغذاء، VO₂ max، النوم، ساعات الوقوف، الوزن، أكسجين الدم، التمارين، وActivity Summary. الكتابة تشمل الماء وبعض مقاييس القلب والمسافة والوزن والتمارين وفق المسار والصلاحية.
        """,
    )
    callout(
        doc,
        "قاعدة الصفر",
        "غياب الصلاحية أو غياب العينة لا يعني قيمة صفر صحيًا. الواجهة والكابتن يجب أن يفرقا بين «0» و«لا توجد بيانات»، خصوصًا النوم والنبض وVO₂ max.",
        "gold",
    )
    heading(doc, "الترابط الخلفي والويدجت", 2)
    bullets(
        doc,
        [
            "مراقبو HealthKit ينعشون البيانات بعد الإعداد الكامل، مع تهدئة للتحديث كي لا يُستنزف الجهاز.",
            "لقطة الويدجت لا تعاد كتابتها أسرع من نحو 20 ثانية، ومراقب اليوم مهدّأ قرابة 60 ثانية.",
            "Hydration Widget يكتب أحداث +250 مل في App Group، ثم يسحبها التطبيق ويدمجها في HealthKit.",
            "البيانات المعروضة تاريخيًا تعتمد على صلاحية HealthKit وتوفر الجهاز/الساعة، لا على خادم AiQo.",
        ],
        bullet_num,
    )

    heading(doc, "الماء الذكي", 1)
    paras(
        doc,
        """
        الهدف الافتراضي 2500 مل. عندما يتوفر الوزن، يوصي النظام بـ 32.5 مل لكل كغ، ثم يقيد الناتج بين 1500 و4000 مل ويقربه إلى أقرب 100 مل. المستخدم يستطيع التسجيل من التطبيق أو الويدجت، وتصبح HealthKit مرجعًا مشتركًا مع بقية منظومة Apple.

        نافذة اليقظة الافتراضية 08:00–22:00، والهدوء 22:00–07:00، وفاصل التذكير 25 دقيقة. يحسب المحرك كمية متوقعة خطيًا حسب وقت اليوم ثم يصنف الوتيرة: متقدم عند 110% فأكثر، على المسار عند 90%، متأخر عند 60%، ومتأخر جدًا تحت ذلك. بعد التأخر يمكن إنشاء تذكير لطيف بعد 30 دقيقة أو أقوى بعد 10 دقائق عبر NotificationBrain.
        """,
    )
    flow(
        doc,
        "الكمية الفعلية ÷ الكمية المتوقعة الآن ← تصنيف الوتيرة ← مراعاة الهدوء والتبريد والميزانية ← تذكير أو صمت",
    )

    heading(doc, "النوم: من العينات إلى قرار", 1)
    paras(
        doc,
        """
        طبقة النوم تقرأ مراحل HealthKit: مستيقظ، REM، Core، وDeep. تجمع العينات المتجاورة، تدمج المصادر بعناية، وتحافظ على دقة المصدر حين توجد بيانات متعددة. تُحسب درجة تقريبية من 0 إلى 100 اعتمادًا على المدة، توقيت النوم، والانقطاعات، وتُعرض الاتجاهات عبر يوم وأسبوع وشهر وسنة وكل الوقت.

        Smart Wake يبني نافذة استيقاظ حول وقت محدد، ويقترح نافذة 10 أو 20 أو 30 دقيقة مع خيار مميز وثلاثة بدائل، مستفيدًا من دورات نوم تقريبية. AlarmKit يدير منبهًا واحدًا تحت سيطرة التطبيق بعد طلب التفويض. هذه ميزة جدولة؛ لا تدّعي أن الهاتف يقيس المرحلة لحظيًا كجهاز مختبر نوم.
        """,
    )
    heading(doc, "وكيل النوم الذكي", 2)
    bullets(
        doc,
        [
            "المسار المفضل محلي عبر Apple Intelligence، ويستلم ملخصًا محسوبًا لا صورًا أو ملفًا خامًا.",
            "مقيّم الجودة يرفض الناتج القصير أو غير العملي؛ يتوقع أكثر من نحو 55 حرفًا، جملتين، إجراءً، وذكر مقياس.",
            "عينات مراحل النوم الخام لا تغادر الجهاز في التصميم الحالي.",
            "عند تعذر النموذج المحلي أو ضعف الجودة قد يرسل المسار ملخص نوم مجمعًا إلى السحابة بعد الموافقة؛ لذا عبارة «لا شيء عن النوم يغادر الجهاز» أوسع من الواقع.",
            "يوجد fallback حسابي ثابت حتى لا تكون الشاشة فارغة عند فشل الذكاء.",
        ],
        bullet_num,
    )
    heading(doc, "المراقبة والتنبيهات", 2)
    paras(
        doc,
        """
        anchored query للنوم يسمح بالتقاط نتائج جديدة بعد استيقاظ المستخدم. يمكن عندها تحديث البطاقة وإطلاق منطق الكابتن أو الدين؛ لكن التنبيهات العميقة تمر عبر بوابات الاشتراك والميزانيات والهدوء. لا ينبغي إنشاء تنبيه طبي قطعي من ليلة واحدة؛ النظام مصمم للتوجيه والاتجاه، لا للتشخيص.
        """,
    )


def add_gym(doc: Document, bullet_num: int, decimal_num: int) -> None:
    part(
        doc,
        "الرابع",
        "النادي: التمرين، الخطة، المعركة، والإنجاز",
        "تفكيك أوسع وحدة في المنتج: من بدء الجلسة الحية إلى الجري الخارجي والخطط المتعددة الأسابيع والتحديات.",
    )
    heading(doc, "بنية النادي", 1)
    table(
        doc,
        ["القسم الداخلي", "الغرض", "الحالة"],
        [
            ["Body", "اختيار نشاط وبدء جلسة حية.", "مفعّل"],
            ["Plan", "إنشاء خطة شخصية وتشغيل مهامها.", "مفعّل؛ العمق حسب الاشتراك"],
            ["Peaks", "مطاردة إنجازات عالمية بخطة دورية.", "عرض Max؛ تنفيذ كامل Pro/Trial"],
            ["Battle", "رحلة مراحل ومهام قابلة للترقية.", "مفعّل"],
            ["Impact", "ملخص أثر وإنجازات المستخدم.", "مفعّل"],
        ],
        [1560, 4680, 3120],
    )

    heading(doc, "Body ومكتبة الأنشطة", 1)
    paras(
        doc,
        """
        يحتوي الكود على 23 نموذج نشاط/تمرين، مرتبة بصريًا في فئات، وتغطي: Cinematic Grind، Cardio/Zone 2، الجري الخارجي، الجري، المشي، الدراجات، السباحة، القوة، HIIT، اليوغا، الفروسية، الكاليستنكس، البيلاتس، الامتنان، الدراجة الداخلية، Elliptical، Stair Stepper، كرة القدم، البادل، السلة، الملاكمة، الفنون القتالية، ونط الحبل.

        ليست كل بطاقة مجرد مؤقت. LiveWorkoutSession يدير الزمن، النبض، السعرات، المسافة، الاتصال بالساعة، التوقف والاستئناف، وحالة الجلسة. بعض الأنشطة تستفيد من HealthKit workout session على Apple Watch، بينما يضيف الهاتف Live Activity وDynamic Island ومزامنة الملخص.
        """,
    )
    heading(doc, "Zone 2 دون استخدام اليدين", 2)
    bullets(
        doc,
        [
            "يتابع النبض مقابل النطاق المستهدف ويحوّل الحالة إلى توجيه صوتي.",
            "يستخدم Apple TTS للمسار الفوري، ويحتفظ بسجل زمني للجلسة.",
            "لا ينبغي أن يضمن دقة النطاق دون عمر وبيانات ملائمة؛ هو تقدير تدريبي، لا قياس سريري.",
            "بعد النهاية يكتب تمرين HealthKit، يحفظ السجل، ويستطيع تشغيل توجيه afterWorkout للمقارنة.",
        ],
        bullet_num,
    )

    heading(doc, "الجري الخارجي", 1)
    paras(
        doc,
        """
        Outdoor Run يستخدم الموقع لرسم المسار، يحفظ حالة الجلسة النشطة حتى لا تضيع عند الانتقال للخلفية، ويولد ملخصًا وصورة مسار وإعادة عرض. background location موجود ضمن أوضاع التطبيق، لذلك يجب أن يظل الطلب سياقيًا وواضحًا: لا حاجة للموقع الدائم عند عدم وجود جلسة.

        المسافة والسرعة والنبض والسعرات تأتي من مصادر مختلفة وقد تتأخر في المزامنة. التصميم يجب أن يعرض دلالة «قيد التحديث» عند الحاجة بدل تبديل أرقام حاد، خصوصًا أثناء انتقال القيادة بين الهاتف والساعة.
        """,
    )

    heading(doc, "Plan: الخطة الشخصية", 1)
    numbered(
        doc,
        [
            "يدخل المستخدم الهدف والخبرة والأيام والوقت والرياضة والقيود والإصابات.",
            "اختياريًا يرفق صورة جسم بعد موافقة منفصلة؛ الصورة ليست إذنًا عامًا لكل تحليل صوري.",
            "تُنشأ خطة يومية مع أسابيع ومهام وتمارين وحجم وشدة.",
            "يشغل المستخدم اليوم المحدد، وتُسجل النتائج في WorkoutHistory.",
            "يجمع النظام إحصاءات أسبوعية، ويمكن للكابتن إنشاء تفسير أو تعديل أعمق.",
        ],
        decimal_num,
    )
    paras(
        doc,
        """
        واجهة الإدخال تعرض مددًا تشمل 1 و2 و4 و8 أسابيع، لكن بوابات التنفيذ ليست متطابقة تمامًا. TierGate يسمح بحد عملي 1 أسبوع لـ Max و4 أسابيع لـ Pro، وخدمة السحابة تعامل أي خطة تتجاوز أسبوعًا كميزة Pro. وجود خيار 8 أسابيع في واجهة أو نموذج لا يعني أنه متاح فعليًا لكل مستوى؛ هذه فجوة يلزم توحيدها.

        في حال فشل السحابة توجد بدائل ثابتة حتى لا يعلق المستخدم. لكن الخطة الصحية الحساسة يجب ألا تتجاهل إشارات الحمل أو القلب أو الضغط أو الجراحة الحديثة التي جُمعت في HealthScreeningStore. هذه البيانات تُحقن كقيود إلزامية في الكابتن وفي مولدات الخطة حيث يوجد تكامل.
        """,
    )

    heading(doc, "Battle وQuestKit", 1)
    paras(
        doc,
        """
        QuestKit يبني رحلة من 10 مراحل، في كل مرحلة 5 مهام حالية. لا تُفتح المرحلة التالية إلا بعد وصول كل مهام السابقة إلى Tier 3. أنواع المهمة: يومية، أسبوعية، مرة واحدة، سلسلة، تراكمية، أو مركبة. مصادر الإثبات: يدوي، ماء، HealthKit، كاميرا، مؤقت، تمرين، اجتماعي، مطبخ، مشاركة، أو تعلم.

        Stage 2 Slot 3 يختار عبر الأعلام بين Learning Spark Stage 2 أو Plank Ladder أو بديل. الإعداد الحالي يفعّل Learning Spark ويغلق Plank Ladder. التحقق بالكاميرا له حد أعلى معلن 95%، بينما شهادات التعلم يمكن التحقق منها محليًا بـ Vision OCR وFoundation Models بعد الموافقة، مع بقاء الصورة على الجهاز وإمكانية حفظ URL مرجعي فقط.
        """,
    )
    table(
        doc,
        ["المكوّن", "القاعدة التنفيذية"],
        [
            ["عدد المراحل", "10"],
            ["المهام في المرحلة", "5"],
            ["فتح المرحلة", "كل مهام السابقة Tier 3"],
            ["Learning Spark 1", "1000 XP عند الاستحقاق الصريح"],
            ["Learning Spark 2", "2000 XP عند الاستحقاق الصريح"],
            ["دقة الكاميرا", "لا تتجاوز 95% في العرض"],
            ["تغيير الأعلام", "يحتاج إعادة تشغيل كي يعاد بناء تعريفات المهمة"],
        ],
        [3120, 6240],
    )
    callout(
        doc,
        "مهم",
        "ليست كل مهمة في QuestKit تمنح XP. الجائزة تُطبق فقط عندما يحدد التعريف قيمة XP صريحة؛ العرض الذي يوحي بجائزة عامة لكل تقدم سيكون مضللًا.",
        "gold",
    )

    heading(doc, "XP والمستويات والسلسلة", 1)
    paras(
        doc,
        """
        LevelStore هو المرجع الحالي للمستوى، وفيه 50 مستوى بعتبات تبدأ من 0 وتصل إلى 1,183,000 XP. دروع التقدم: خشب 1–4، برونز 5–9، فضة 10–14، ذهب 15–19، بلاتين 20–24، ألماس 25–29، Obsidian 30–34، وLegendary من 35. توجد LevelSystem أخرى بسبع رتب، وهي ازدواج يحتاج إزالة أو جعلها واجهة واحدة.

        XP التمرين يحسب أساسًا من السعرات + الدقائق + مجموع أرقام تقدير إجمالي ضربات القلب. العملات تضيف 1 لكل 100 خطوة، و1 لكل 50 سعرة، و2 لكل دقيقة إذا تجاوز متوسط النبض 115. هذه معادلات لعب، وليست تقديرًا علميًا لجودة التدريب.

        السلسلة تعتبر اليوم نشطًا عند تحقيق أحد الشروط المنفذة مثل 5000 خطوة أو تمرين أو مدة نشاط كافية. هناك حالة خطر قرب 22 ساعة دون إنجاز، وسجل يصل 90 يومًا مع تركيز على آخر 30 يومًا ومؤشر اتساق سبعة أيام.
        """,
    )

    heading(doc, "Legendary Peaks", 1)
    paras(
        doc,
        """
        Peaks يحوّل أرقامًا أسطورية إلى برامج طويلة مقسمة إلى Foundation وBuild وIntensify وDeload كل أربعة أسابيع وTaper أخير. في الكود ثمانية أهداف seed: 152 ضغطًا في دقيقة خلال 16 أسبوعًا، بلانك 9.5 ساعات خلال 24 أسبوعًا، 70 سكوات في دقيقة خلال 10 أسابيع، 228.93 كم مشيًا في 24 ساعة خلال 20 أسبوعًا، 48 Burpee في دقيقة خلال 12 أسبوعًا، 62 Pull-up في دقيقة خلال 16 أسبوعًا، حبس نفس 24.37 دقيقة خلال 12 أسبوعًا، و210 ألف خطوة في 24 ساعة خلال 16 أسبوعًا.

        يوجد HRR Engine لاختبار التعافي عبر Apple Watch، ومراجعة أسبوعية يمكن أن تحصل على سرد سحابي. Max يستطيع رؤية المحتوى، بينما Pro/Trial يفتح المسار الكامل. الأهداف شديدة ويجب أن تبقى محاطة بتحذيرات وفحص ملاءمة؛ بعضها غير مناسب لمعظم المستخدمين دون إشراف.
        """,
    )


def add_kitchen_vibe_kernel(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "الخامس",
        "المطبخ، My Vibe، وKernel",
        "ثلاث وحدات تربط التغذية والحالة الذهنية والانتباه الرقمي بالفعل اليومي.",
    )
    heading(doc, "المطبخ: من وجبة اليوم إلى قائمة التسوق", 1)
    paras(
        doc,
        """
        KitchenScreen يعرض الإفطار والغداء والعشاء، متتبع السعرات والبروتين والكربوهيدرات والدهون والألياف، الثلاجة، وخطة الوجبات. يحتوي المشروع على 18 وجبة مدمجة في meals_data.json كي تبقى الشاشة مفيدة حتى قبل أي توليد سحابي.

        HomeKitchenRootView يستخدم AccessManager، ولذلك يرى المستخدم المجاني Paywall عند فتح تبويب المطبخ رغم أن التبويب ظاهر. هذه إحدى نقاط الالتباس: الظهور لا يساوي الإتاحة، وبوابة المطبخ تعتمد طبقة اشتراك أقدم من TierGate.
        """,
    )
    heading(doc, "الثلاجة الذكية", 2)
    flow(
        doc,
        "التقاط صورة ← موافقة الصورة ← إزالة EXIF/GPS وتصغيرها ← Gemini Vision ← عناصر موحّدة ← مخزون ← خطة وتسوق",
    )
    bullets(
        doc,
        [
            "الكاميرا تستخدم AVCapture، والتحليل يحتاج موافقة وصلاحية مستوى مدفوع.",
            "الصورة السحابية تقيد إلى 1280 بكسل تقريبًا وتضغط JPEG بجودة 0.78 بعد حذف metadata.",
            "الناتج يطابق 69 IngredientKey قياسيًا ضمن بروتين، كربوهيدرات، خضار، فواكه، ألبان، دهون، مشروبات، وأخرى.",
            "المخزون والخطة المثبتة وقائمة التسوق محفوظة محليًا كـ JSON في UserDefaults.",
            "بعض مسارات الكاميرا تفحص captainChat/Max وبعضها يفحص photoAnalysis/Pro؛ يلزم توحيد قرار المنتج.",
        ],
        bullet_num,
    )
    heading(doc, "مولد خطة الطعام", 2)
    paras(
        doc,
        """
        الخطة تقبل 3 أو 7 أيام في المسار الحالي. كل وجبة تحمل سعرات وماكروز وألياف ومكونات، ثم يُشتق منها ما ينقص المستخدم وقائمة التسوق وبدائل المكونات. عند غياب السحابة يوجد مولد deterministic بدل شاشة خطأ.

        الخطة الأسبوعية تساوي أسبوعًا واحدًا، ولذلك تمر عادة ضمن Max، بينما الصور أو التحليل الأعمق قد يرفعان المتطلب. يجب ألا يُفترض أن أي خطة غذائية علاجية؛ الحساسية، الأمراض المزمنة، والحمل تحتاج رسائل سلامة واضحة وقيودًا قبل التوليد.
        """,
    )

    heading(doc, "My Vibe", 1)
    table(
        doc,
        ["الحالة", "الوقت الافتراضي", "الصوت المحلي"],
        [
            ["Awakening", "05:00–09:00", "Serotonin Flow"],
            ["Deep Focus", "09:00–12:00", "Gamma Flow"],
            ["Peak Energy", "12:00–17:00", "Sound of Energy"],
            ["Recovery", "17:00–21:00", "Hypnagogic"],
            ["Ego Death", "21:00–05:00", "Theta Trance"],
        ],
        [2340, 2340, 4680],
    )
    paras(
        doc,
        """
        VibeOrchestrator يراجع الحالة كل 30 ثانية تقريبًا ويعمل crossfade بين ملفات صوت محلية، مع إمكانية أن يتجاوز DJ الاختيار. Spotify مدعوم عبر App Remote وWeb API مع PKCE، وصلاحيات قراءة الأكثر استماعًا وحالة التشغيل وتعديل التشغيل. الرموز تحفظ في Keychain مع مسار هجرة.

        My Vibe حاليًا خلف AccessManager كميزة مدفوعة. الاسم «Ego Death» قوي تسويقيًا وقد لا يناسب كل سياق ثقافي أو صحي؛ يحتاج تعريفًا هادئًا داخل الواجهة حتى لا يُفهم كمفهوم علاجي أو نفسي.
        """,
    )

    heading(doc, "Kernel: تحويل المشتت إلى حركة", 1)
    paras(
        doc,
        """
        Kernel يستخدم FamilyControls وManagedSettings وDeviceActivity لانتقاء التطبيقات أو الفئات أو النطاقات ثم حجبها. يوجد تطبيق رئيسي وثلاثة امتدادات: مراقب Device Activity، إعداد شاشة الحجب، وإجراء الحجب. الاختيارات والتحديات تمر عبر App Group بصيغة JSON كي تعمل الامتدادات دون فتح التطبيق.

        التنفيذ الحالي يتيح Kernel لكل المستويات: المجاني يختار تطبيقًا/فئة/نطاقًا واحدًا، والمدفوع عددًا غير محدود. هذا يخالف تعليقات أو مواد أقدم تصفه Max-only. يوجد نمطان: Smart ينتظر شرط الاستخدام والخمول، وHard يطبق الحجب مباشرة.
        """,
    )
    heading(doc, "سلم التصعيد", 2)
    table(
        doc,
        ["المحاولة", "دقائق قبل التحدي", "خطوات", "جلسة", "طاقة"],
        [
            ["1", "15", "40", "5 دقائق", "20 kcal"],
            ["2", "8", "90", "4 دقائق", "60 kcal"],
            ["3", "5", "160", "3 دقائق", "150 kcal"],
            ["4", "3", "260", "دقيقتان", "350 kcal"],
            ["5", "2", "400", "1.5 دقيقة", "700 kcal"],
            ["بعد 5", "—", "+250 كل مرة حتى 2500", "1.5 دقيقة", "لا تجاوز بالطاقة"],
        ],
        [1260, 1800, 2340, 1800, 2160],
    )
    paras(
        doc,
        """
        Bio Score يرفع الصعوبة 10% لكل نقطة مستنتجة من الخمول أو التوتر أو الوقت المتأخر. توجد أنواع تحدي تنفس وهدوء نبض في النماذج، لكن المولد الحالي يصدر تحديات خطوات. يستطيع المستخدم تعطيل Kernel عبر مسار متعمد الاحتكاك؛ يجب أن يبقى هناك مخرج واضح لأن منع الوصول إلى التطبيقات حساس وظيفيًا.

        أسماء التطبيقات والرموز المختارة تبقى محلية في منظومة FamilyControls/App Group. لا يرى AiQo قائمة التطبيقات كنصوص عادية بالطريقة التي يراها تطبيق تحليلات تقليدي، ويجب ألا يحاول رفع tokens إلى الخادم.
        """,
    )
    callout(
        doc,
        "مبدأ تصميم",
        "Kernel ناجح عندما يجعل العودة إلى التطبيق المشتت قرارًا واعيًا، لا عندما يعاقب المستخدم. التصعيد يجب أن يتوقف عند حدود السلامة وإمكانية الوصول والإرهاق.",
        "mint",
    )


def add_captain(doc: Document, bullet_num: int, decimal_num: int) -> None:
    part(
        doc,
        "السادس",
        "الكابتن: العقل الحواري والذاكرة",
        "شرح مسار الرسالة، النماذج المحلية والسحابية، بناء السياق، الذاكرة، التوجيهات، الصوت، والسلامة.",
    )
    heading(doc, "الكابتن كمنتج داخل المنتج", 1)
    paras(
        doc,
        """
        الكابتن هو الواجهة التي توحد الوحدات. المستخدم لا يحتاج معرفة أين توجد بيانات النوم أو خطة التمرين؛ يستطيع سؤال شخصية واحدة. خلف هذه البساطة توجد طبقات sensing وmemory وreasoning وinference وprivacy وproactive وlearning وpersona وwellbeing وobservability وdirectives.

        IntentClassifier يصنف الرسالة إلى عام، تمرين، تغذية، نوم، تحدي، Vibe، دعم عاطفي، أو تعافٍ، ويكشف الأزمة كمسار مستقل. التصنيف يغير وزن الاسترجاع، سياق الشاشة، والأداة أو النموذج المختار. الرد قد يكون نصًا فقط أو استجابة منظمة فيها Quick Replies، خطة تمرين، خطة وجبة، اقتراح Spotify، ذاكرة محفوظة، أو تذكير.
        """,
    )

    heading(doc, "مسارا المجاني والمدفوع", 1)
    table(
        doc,
        ["البعد", "Free على الجهاز", "Max / Pro"],
        [
            ["النموذج", "Apple Foundation Models", "Gemini 2.5 Flash؛ Preview مغلق افتراضيًا"],
            ["مكان التنفيذ", "على الجهاز", "سحابي بعد موافقة، مع fallbacks"],
            ["السياق", "جلسة LanguageModelSession الحالية", "نافذة محادثة + digest + ذاكرة مسترجعة"],
            ["الذاكرة الدائمة", "لا", "نعم وفق المستوى"],
            ["البيانات الحية", "ملخص HealthKit اليوم", "Bio state وسياق أوسع"],
            ["تخصيص النبرة", "صوت ثابت واسم فقط", "Presets؛ وCustom Style في Pro"],
            ["الخطط الكبيرة", "نصيحة قصيرة ثم دعوة Max", "خطط وبطاقات وتتبع"],
            ["السقف اليومي", "غير محدد بعدد رسائل في الكود", "محكوم بالخدمة والاشتراك لا بقالب ثابت"],
        ],
        [1870, 3430, 4060],
        font_size=8.5,
    )
    paras(
        doc,
        """
        المجاني يعيد استخدام جلسة واحدة كي يتذكر الحوار الحالي، ويعيد ضبطها عند محادثة جديدة أو تعذر السياق. يستدعي الخطوات والنبض والسعرات والنوم والماء، يمرر الناتج على OnDeviceReplySanitizer، يفرض اللهجة العراقية، ثم يصحح الأرقام عبر CaptainFactGuard. الرد مقيد تقريبًا بـ 200 token ونبرة قصيرة. عند طلب برنامج ضخم يعطي «تذوقًا» مفيدًا ثم يوضح أن الخطة المتتبعة والدائمة في Max.

        إذا لم تتوفر Apple Intelligence أو اللغة/المنطقة غير مدعومة، يعرض fallback دافئًا لإعادة المحاولة. هذه ليست بديلًا كاملًا على أجهزة غير مؤهلة؛ لذلك يجب أن يوضح المنتج متطلبات الجهاز دون اتهام الشبكة.
        """,
    )

    heading(doc, "رحلة الرسالة المدفوعة", 1)
    numbered(
        doc,
        [
            "تنظيف النص والتأكد من عدم وجود طلب مكرر قيد التنفيذ.",
            "تصنيف النية وكشف الأزمة؛ الأزمة تتجاوز بعض بوابات الاشتراك حتى لا تُحجب السلامة.",
            "فرض موافقة السحابة إلا لمسار النوم المحلي أو الاستثناءات الآمنة.",
            "حفظ رسالة المستخدم وتحديث سجل المحادثة وخيط السياق.",
            "اكتشاف Directive صريح مثل «بعد كل تمرين قارن وبلغني» وحفظه قبل الرد.",
            "بناء PromptContext من الشاشة والهوية والصحة والتخصيص والذاكرة.",
            "استرجاع حقائق وحلقات وأنماط ومشاعر وعلاقات حسب الميزانية والنية.",
            "تشغيل Safety/Fact/Privacy gates ثم اختيار المحلي أو السحابي.",
            "تحليل الاستجابة المنظمة، تصحيح الأرقام، عرض النص والبطاقات، وحفظ الآثار الموافق عليها.",
            "تسجيل قياسات الجودة والتغذية الراجعة دون تحويل كل المحادثة إلى حدث تحليلي.",
        ],
        decimal_num,
    )
    flow(
        doc,
        "رسالة ← نية/أزمة ← موافقة ← سياق وذاكرة ← PrivacySanitizer ← نموذج ← Parser ← Fact Guard ← واجهة + ذاكرة/تذكير",
    )

    heading(doc, "بناء الـ Prompt", 1)
    paras(
        doc,
        """
        بعض الوثائق تسمي التصميم «سبع طبقات»، لكن PromptComposer الحالي يجمع نحو 15 كتلة وظيفية. الرقم ليس مهمًا بقدر الترتيب: التعليمات العليا والسلامة قبل الشخصية، والحقائق قبل الزخرفة، وعقد المخرجات في النهاية.
        """,
    )
    table(
        doc,
        ["الترتيب التقريبي", "الكتلة", "الهدف"],
        [
            ["1", "Language Lock", "منع تبديل اللغة أو خلطها دون طلب."],
            ["2", "Safety", "حدود طبية وأزمة وإحالة."],
            ["3", "Identity", "من هو الكابتن وكيف يتكلم."],
            ["4", "Stable Profile", "حقائق المستخدم الثابتة المسموح بها."],
            ["5", "Injury Constraints", "قيود الفحص الصحي والإصابات."],
            ["6", "Working Memory", "المعلومات الأقرب للرسالة."],
            ["7", "Conversation State", "ملخص الخيط والالتزامات والتصحيحات."],
            ["8", "Coaching Thesis", "التوجه التدريبي المتماسك."],
            ["9", "Bio State", "خطوات ونوم وتعافٍ وإشارات اليوم."],
            ["10", "Kernel Status", "الانتباه والحجب والتحدي الحالي."],
            ["11", "Circadian Tone", "وقت اليوم والنبرة الملائمة."],
            ["12", "App Knowledge", "قدرات AiQo وحدوده."],
            ["13", "Screen Context", "المكان الذي أتى منه السؤال."],
            ["14", "Medical Disclaimer", "تذكير الحدود عند الحاجة."],
            ["15", "Output Contract", "JSON أو بنية الاستجابة المطلوبة."],
        ],
        [1240, 2800, 5320],
        font_size=8.3,
    )

    heading(doc, "السياسة النموذجية والـ fallback", 2)
    bullets(
        doc,
        [
            "النموذج المستقر هو gemini-2.5-flash.",
            "gemini-3-flash-preview لا يستخدم إلا إذا كان GEMINI_3_PREVIEW مفعّلًا؛ الإعداد الحالي false.",
            "النوم يبدأ محليًا، وبعض أسئلة النوم الصارمة تُعترض كي لا تدخل المسار العام.",
            "عند تعذر السحابة ينتقل النظام إلى محلي أو رد محسوب/مترجم بحسب الوحدة.",
            "CAPTAIN_REAL_STREAMING مغلق افتراضيًا، لذلك لا ينبغي افتراض SSE حقيقي في كل شاشة رغم دعم الوكيل له.",
        ],
        bullet_num,
    )

    heading(doc, "الخصوصية قبل السحابة", 1)
    paras(
        doc,
        """
        PrivacySanitizer يحذف البريد وأرقام الهاتف وUUID وmentions وURLs والأرقام الطويلة وعناوين IP والسلاسل الشبيهة بـ Base64 ومفاتيح API وBearer tokens. يحد المحادثة إلى آخر 16 رسالة تقريبًا أو 6000 حرف، ويمنع إرسال قياسات صحية دقيقة غير لازمة. الاسم الأول والعمر والجنس قد تدخل CloudSafeProfile فقط بعد الموافقة، بينما الارتفاع والوزن الدقيقان يحجبان افتراضيًا.

        الصور مسموحة لمسارات مطبخ/نادي محددة، تُزال منها EXIF وGPS، تُصغّر إلى حد أقصى 1280 بكسل وتضغط. الأخطاء المرسلة إلى Crashlytics تمر أيضًا بمصحح خصوصية. هذه آليات قوية، لكنها لا تجعل النظام «تشفيرًا طرفًا لطرف»؛ مزود السحابة يرى الحمولة المنقحة اللازمة للمعالجة.
        """,
    )
    callout(
        doc,
        "مبدأ الدقة",
        "CaptainFactGuard يعمل بعد التوليد لتصحيح أرقام الخطوات والسعرات والنبض عندما يناقض النموذج المصدر المحلي. لا يعفي ذلك الـ prompt من grounding؛ هو خط دفاع أخير.",
        "mint",
    )

    heading(doc, "الذاكرة: خمسة مخازن، لا صندوق واحد", 1)
    table(
        doc,
        ["المخزن", "ما يحتفظ به", "مثال"],
        [
            ["Semantic", "حقائق مستقرة مع ثقة وأهمية ومصدر.", "يفضل المشي صباحًا."],
            ["Episodic", "مقاطع أو تبادلات ذات حدث.", "أكمل أول 5 كم يوم الجمعة."],
            ["Procedural", "أنماط عمل متكررة.", "ينجح عندما تكون الخطة 20 دقيقة."],
            ["Emotional", "لقطات مزاج وسياق.", "توتر قبل العمل ثم تحسن بعد المشي."],
            ["Relationship", "أشخاص وعلاقات مذكورة بأمان.", "يتدرب مع صديق."],
        ],
        [1870, 3740, 3750],
    )
    paras(
        doc,
        """
        الاسترجاع يوزع ميزانيته افتراضيًا: 40% حقائق، 25% حلقات، 15% أنماط، 10% مشاعر، 10% علاقات. الترتيب يمزج embeddings والكلمات والثقة والحداثة. توجد كذلك مقاييس أسبوعية، consolidation، تأمل شهري، توجيهات، coaching thesis، وconversation threads.

        أرقام الذاكرة ليست واحدة في كل طبقة: SubscriptionTier يعرض 100 للFree و500 لـMax و1000 لـTrial/Pro؛ TierGate يضع سقوفًا دلالية 120 و600 و1200؛ بينما عمق الاسترجاع النشط 0 للمجاني، 18 لـMax، و40 لـPro. المعنى العملي: «سعة التخزين» و«عدد العناصر المسترجعة في كل رد» مختلفان، لكن الواجهة والوثائق تحتاج مصطلحات موحدة.
        """,
    )
    heading(doc, "ضغط المحادثة", 2)
    paras(
        doc,
        """
        الشاشة تحتفظ بنحو 80 رسالة في الذاكرة، لكن النافذة الحية تُقيد إلى 24 رسالة أو نحو 9000 حرف. ConversationCompactor لا يستخدم نموذجًا للتلخيص؛ يبني digest حتميًا يحتفظ بنية البداية، حتى 8 نقاط للمستخدم، 6 التزامات، 4 تصحيحات، وآخر تبادل. تظهر علامة ناعمة في واجهة الخيط عند موضع الضغط حتى لا يكون الاختفاء صامتًا.
        """,
    )

    heading(doc, "التوجيهات القابلة للتنفيذ", 1)
    paras(
        doc,
        """
        إذا قال المستخدم: «بعد كل تمرين حلله وقارنه بالسابق وبلغني»، DirectiveLearner يبحث عن تكرار + فعل + trigger معروف. يُحفظ التوجيه في SwiftData V5 ويُعكس إلى الذاكرة ليستطيع الكابتن تأكيده في الرد نفسه. الحد الأقصى 40 توجيهًا، والتكرار لنفس trigger/action يحدّث الموجود بدل إنشاء نسخ.
        """,
    )
    table(
        doc,
        ["Trigger", "قابل للتنفيذ آليًا الآن", "الفعل"],
        [
            ["afterWorkout", "نعم", "تحليل ومقارنة deterministic ثم إشعار."],
            ["beforeBedtime", "نعم", "إشعار متكرر قبل النوم."],
            ["everyMorning", "نعم", "إشعار صباحي متكرر."],
            ["afterPoorSleep", "معروف لكن غير آلي كاملًا", "يحتاج ربطًا تنفيذيًا."],
            ["weeklyReview", "معروف لكن غير آلي كاملًا", "يحتاج ربطًا تنفيذيًا."],
        ],
        [2340, 3120, 3900],
    )

    heading(doc, "الاستباق والإشعارات", 1)
    paras(
        doc,
        """
        NotificationBrain يعرف نوايا صباحية، دين نوم، خمول، رقم شخصي، تعافٍ، خطر سلسلة، حفظ سلسلة، انسحاب أو زخم، ذاكرة، عاطفة، مزاج، علاقات، أسبوع وشهر، رمضان وعيد وجمعة، إيقاع يومي وطقس، تجربة اشتراك، إنجاز، تمرين، وترطيب.

        قبل الجدولة يطبق quiet hours وcooldown وglobal budget وdedupe وبوابة المستوى. توجد قيم مختلفة بين الطبقات: NotificationBrain يستخدم تقريبًا 3 يوميًا و4 ساعات للأساسي، 5 و3 ساعات لـMax، و6 وساعتين لـPro؛ بينما SubscriptionTier/TierGate يعرضان أرقامًا أخرى مثل 2/4/7 أو 0/4/7. هذه ليست مجرد وثيقة قديمة؛ هي سياسة موزعة يجب توحيدها كي لا يتلقى المستخدم أكثر أو أقل من الوعد.
        """,
    )

    heading(doc, "الصوت", 1)
    bullets(
        doc,
        [
            "Apple TTS يوفر الصوت الفوري المحلي ويرافق الجلسات والإرشاد.",
            "المسار المميز يستخدم MiniMax عبر captain-voice بعد موافقة الصوت وإعداد المفاتيح.",
            "يوجد Voice Cache لتقليل الاستدعاءات وإعادة التشغيل.",
            "TierGate يصنف premiumVoice كميزة Pro، لكن بعض واجهات الاختيار أو وثائق أقدم تربطه بالمدفوع عمومًا؛ يجب أن تكون شاشة الشراء هي المرجع المرئي.",
            "CAPTAIN_VOICE_CLOUD مفعّل في Info.plist، لكن نجاحه الفعلي يعتمد على إعداد Supabase والمزود.",
        ],
        bullet_num,
    )

    heading(doc, "السلامة النفسية والطبية", 1)
    paras(
        doc,
        """
        CrisisDetector يسبق التوليد العام. الاستجابة تتدرج بين check-in لطيف، انعكاس، وإحالة مهنية، مع موارد بحسب الإمارات والسعودية والعراق ومسار عالمي. رسائل الأزمة لا تُمنع بسبب الاشتراك. يجب أن تظل الإحالة محدثة جغرافيًا، وألا يدعي الكابتن أنه جهة طوارئ أو معالج.

        HealthScreeningStore يضع الحمل ومشاكل القلب والضغط والجراحة الحديثة داخل القيود. هذه الإشارات لا تسمح للكابتن بالتشخيص؛ وظيفتها تقليل المخاطرة، تخفيف الشدة، وطلب مراجعة مختص عند الحاجة.
        """,
    )

    heading(doc, "الرصد والتعلم", 2)
    bullets(
        doc,
        [
            "BrainBus يوحّد أحداثًا داخلية بين الاستشعار والذاكرة والتنبيهات.",
            "التعلم الأسبوعي يدمج feedback والالتزامات والاتجاهات في thesis أكثر اتساقًا.",
            "لوحات تشخيص ومقاييس داخلية تساعد على تفسير لماذا اختير رد أو ذاكرة.",
            "Memory V4/V5 يستخدم ترحيلًا وshadow-write في بعض المسارات؛ الفشل الصامت يحمي التجربة لكنه قد يخفي فقدًا غير ملحوظ.",
            "إعدادات الذاكرة تسمح للمستخدم بالمراجعة والحذف، وهي إضافة محورية في 1.0.8.",
        ],
        bullet_num,
    )


def add_watch_widgets_reports(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "السابع",
        "Apple Watch، الويدجت، التقارير، والملف",
        "الأسطح التي تجعل AiQo حاضرًا خارج التبويب الرئيسي وتغلق حلقة القياس والمراجعة.",
    )
    heading(doc, "تطبيق Apple Watch", 1)
    paras(
        doc,
        """
        تطبيق الساعة يدعم تسعة أنواع جلسة بارزة: مشي وجري داخلي/خارجي، دراجة، HIIT، قوة، يوغا، وسباحة. WatchWorkoutManager يبني HKWorkoutSession ويعرض النبض والسعرات والمسافة والوقت، مع إيقاف واستئناف وإنهاء. بعد الجلسة ترسل الساعة ملخصًا للهاتف وتستطيع عرض التعافي.

        الاتصال يستخدم WatchConnectivity. أثناء الجلسة تُرسل snapshot متقاربة—يوجد دفع دوري يقارب 0.75 ثانية وتهدئة حول 0.5 ثانية—كي تبقى واجهة الهاتف وLive Activity متزامنتين دون إرسال كل تغير خام. إذا انقطع الاتصال، يجب أن تبقى جلسة HealthKit على الساعة مرجعًا ثم تتم المصالحة لاحقًا.
        """,
    )
    bullets(
        doc,
        [
            "الرئيسية على الساعة تعرض خطوات وسعرات ومسافة ونوم وماء وأهدافًا.",
            "المعالم milestones تنشئ تغذية راجعة أثناء التمرين.",
            "جلسة Watch قد تكون مالكة للتمرين؛ الهاتف ليس دائمًا المصدر النهائي.",
            "اختبارات الساعة موجودة كTargets منفصلة، إضافة إلى UI Tests.",
        ],
        bullet_num,
    )

    heading(doc, "الويدجت وLive Activities", 1)
    table(
        doc,
        ["السطح", "المحتوى", "التفاعل"],
        [
            ["Home Widget", "حلقات/ملخص اليوم بأحجام small وmedium وaccessory.", "فتح التطبيق أو الوجهة."],
            ["Hydration Widget", "الماء والتقدم.", "زر تفاعلي +250 مل عبر App Group."],
            ["Workout Live Activity", "وقت، نبض/zone، حالة الجلسة.", "العودة للجلسة من Lock Screen/Dynamic Island."],
            ["Watch Widgets", "ملخص أسبوعي وحلقات.", "نظرة سريعة على الساعة."],
        ],
        [1870, 4680, 2810],
    )

    heading(doc, "التقرير الأسبوعي", 1)
    paras(
        doc,
        """
        التقرير يقارن آخر سبعة أيام بالسبعة السابقة ويولد Score من 0 إلى 100، مع الخطوات والسعرات والمسافة والنوم والماء والتمارين وسلسلة يومية. يمكن مشاركة صورة التقرير وتصدير CSV وPDF للتقرير الحالي.

        توجد وثائق أقدم تقول إن التطبيق يصدر كل بيانات الصحة بصيغ CSV وJSON وPDF. التنفيذ الحالي لـHealthDataExporter لا يثبت تصدير JSON شاملًا لكل HealthKit؛ القيمة المتحققة هي تقرير أسبوعي وملفاته، لذا يجب تعديل الوعد أو استكمال المصدر الشامل.
        """,
    )

    heading(doc, "الصور التقدمية", 1)
    paras(
        doc,
        """
        صور التقدم تحفظ محليًا في Documents/ProgressPhotos كـJPEG بجودة تقارب 0.85، والبيانات الوصفية في UserDefaults. الشاشة تحمل 20 عنصرًا في الصفحة، وتتيح مقارنة أول وآخر صورة وإظهار تغير الوزن إن توفر.

        هذا يختلف عن وثائق تذكر SwiftData. الصور حساسة جدًا: يجب حماية النسخ الاحتياطية، توضيح ما إذا كانت تدخل iCloud Backup، وعدم إرسالها للكابتن إلا بموافقة منفصلة ومقصودة.
        """,
    )

    heading(doc, "الملف والإعدادات", 1)
    table(
        doc,
        ["المجموعة", "ما يتوفر"],
        [
            ["الهوية", "صورة محلية، اسم وقياسات وهدف."],
            ["التقدم", "مستوى، XP، سلسلة، تقرير أسبوعي، صور."],
            ["الخصوصية", "موافقة AI، موافقة الصوت، صور الجسم، الذاكرة، إخلاء طبي."],
            ["التجربة", "لغة التطبيق، لغة الإشعارات، نبرة الكابتن."],
            ["المنتج", "الاشتراك، Kernel، الإحالات، المساعدة والقانوني."],
            ["الحساب", "تسجيل خروج وحذف الحساب."],
        ],
        [2340, 7020],
    )
    paras(
        doc,
        """
        سحب موافقة AI يضع التطبيق في offline-only للمسارات المتأثرة؛ لا ينبغي أن يعطل القياسات المحلية. موافقة الصوت مستقلة عن موافقة النص، وموافقة صورة الجسم مستقلة عن صورة الثلاجة. الفصل بين الأذونات مبدأ مهم كي لا تتحول موافقة واحدة إلى تفويض واسع.
        """,
    )

    heading(doc, "Tribe وArena", 1)
    paras(
        doc,
        """
        Tribe يحتوي نحو 58 ملفًا ويغطي إنشاء قبيلة والانضمام بكود دعوة، حد خمسة أعضاء، طاقة وSparks ومهمات ولوحات صدارة وتحديات أسبوعية وHall of Fame وقادة إمارات وسجل. SupabaseArenaService يعرف جداول profiles وarena_tribes وmembers وparticipations وweekly_challenges وhall_of_fame_entries.

        رغم اكتمال جزء كبير من البنية، الإعداد الحالي يجعل TRIBE_FEATURE_VISIBLE=false وTRIBE_BACKEND_ENABLED=false. لذلك Tribe ليس ميزة مستخدم حية في النسخة المفحوصة. وجود TRIBE_SUBSCRIPTION_GATE=false لا يفتحها للعامة؛ هو فقط يعطل بوابة الدفع داخل مسار مخفي أصلًا.
        """,
    )
    callout(
        doc,
        "تصنيف صحيح",
        "Tribe = «منفذ وغير ظاهر»، لا «غير موجود» ولا «متاح». هذا الفرق مهم في خارطة الطريق والتقييم الاستثماري ودعم المستخدم.",
        "blue",
    )


def add_architecture(doc: Document, bullet_num: int, decimal_num: int) -> None:
    part(
        doc,
        "الثامن",
        "البنية التقنية ودورة التشغيل",
        "خريطة الأهداف والوحدات، الإقلاع، التخزين، الاعتماديات، والأعلام التي تغير السلوك.",
    )
    heading(doc, "صورة البنية", 1)
    flow(
        doc,
        "SwiftUI Screens ↔ ViewModels/Stores ↔ Domain Engines/Actors ↔ HealthKit & SwiftData & UserDefaults ↔ Supabase/AI/Spotify",
    )
    paras(
        doc,
        """
        التطبيق مكتوب أساسًا بـSwift وSwiftUI، مع استخدام Actors في الخدمات الحساسة للتزامن، و@MainActor لمخازن الواجهة. لا توجد معمارية واحدة مفروضة على كل الملفات؛ الغالب MVVM مع Services وStores وEngines. الوحدات الأكبر—Captain وGym—تستخدم تقسيمًا مجاليا أعمق من الوحدات القديمة.

        حجم المستودع يفسر بعض الازدواج: نحو 698 ملف Swift و68 ملف Markdown عند اللقطة، وأكثر من 155 ملفًا للكابتن و102 للنادي. الجودة هنا تعتمد على تحديد «مصدر حقيقة» لكل مجال؛ المشكلة ليست نقص الكود، بل تعدد طبقات تؤدي الدور نفسه في الاشتراك والمستوى وHealthKit.
        """,
    )

    heading(doc, "Targets وامتدادات Xcode", 1)
    table(
        doc,
        ["الهدف", "الدور"],
        [
            ["AiQo", "تطبيق iPhone الرئيسي."],
            ["AiQoTests / AiQoUITests", "اختبارات الوحدة والواجهة."],
            ["AiQoWatch Watch App", "تطبيق Apple Watch."],
            ["AiQoWatchWidgetExtension", "ويدجت الساعة."],
            ["Watch Tests / UI Tests", "اختبارات الساعة."],
            ["AiQoWidgetExtension", "ويدجت iOS ومنها الترطيب."],
            ["AiQoWorkoutLiveAttributesExtension", "Live Activity وDynamic Island للتمرين."],
            ["KernelDeviceActivityMonitor", "مراقبة وقت النشاط والحجب."],
            ["KernelShieldConfiguration", "تصميم شاشة الحجب."],
            ["KernelShieldAction", "معالجة إجراء المستخدم على الحجب."],
        ],
        [3740, 5620],
    )
    paras(
        doc,
        """
        Deployment Target هو iOS 26.2 وwatchOS 26.2، وSwift language version في المشروع 5.0. Bundle الرئيسي com.mraad500.aiqo. ارتفاع حد النظام متعمد لأنه يتيح Foundation Models وAlarmKit وواجهات أحدث، لكنه يقلل قاعدة الأجهزة المتوافقة ويجعل توفر Apple Intelligence شرطًا منفصلًا عن توفر iOS.
        """,
    )

    heading(doc, "إقلاع التطبيق", 1)
    numbered(
        doc,
        [
            "تهيئة CrashReporting المشروط وCrashReporter المحلي.",
            "تشغيل WatchConnectivity ومراقبة الشبكة والتحليلات المحلية.",
            "قراءة Remote/Info flags وتكوين Localization.",
            "تهيئة PurchaseManager وFreeTrialManager القديم والـ entitlements.",
            "بناء SwiftData للحياة اليومية والتمارين وArena.",
            "بناء حاوية ذاكرة الكابتن: V5 عند MEMORY_V4، وإلا V3، مع migration وfallback آمن أو in-memory.",
            "حقن MemoryStore والمخازن الخمسة وDirectiveStore وThreadManager.",
            "تسجيل قرابة 15 trigger إذا كانت الذاكرة الحديثة مفعلة، وتشغيل BrainBus.",
            "تهيئة NotificationBrain وCrisisDetector والاختصارات.",
            "بعد اكتمال onboarding: تشغيل مراقبي HealthKit والجدولة الخلفية.",
        ],
        decimal_num,
    )
    callout(
        doc,
        "استراتيجية التعافي",
        "إذا فشل مخزن SwiftData الحديث، يحاول التطبيق التراجع إلى مخطط أقدم أو مخزن مؤقت بدل الانهيار. هذا يحمي الفتح لكنه قد يخفي مشكلة ترحيل؛ يجب تسجيل نسبة fallback ومتابعتها.",
        "gold",
    )

    heading(doc, "التخزين المحلي", 1)
    table(
        doc,
        ["التقنية", "البيانات البارزة", "حدودها"],
        [
            ["HealthKit", "نشاط، نوم، قلب، ماء، تمارين، وزن.", "صلاحيات Apple ومصدر مشترك."],
            ["SwiftData", "ذاكرة الكابتن، التوجيهات، سجلات يومية، مهام، Arena.", "مخططات وترحيلات متعددة."],
            ["UserDefaults", "إعدادات، onboarding، ثلاجة، خطط مثبتة، metadata.", "ليس للسر الثقيل أو الصور."],
            ["Keychain", "رموز Spotify وأسرار/حالة حساسة مختارة.", "محلي محمي؛ يحتاج سياسة حذف."],
            ["Documents", "صور التقدم وملفات قابلة للمشاركة.", "قد تدخل النسخ الاحتياطية."],
            ["Application Support", "analytics.jsonl وcrash_log.jsonl.", "محلي، محدود بعدد سجلات."],
            ["App Groups", "ويدجت وKernel snapshots وأوامر.", "مشترك بين الامتدادات فقط."],
        ],
        [1870, 4210, 3280],
        font_size=8.5,
    )

    heading(doc, "الاعتماديات والمنصات", 1)
    bullets(
        doc,
        [
            "Supabase Swift 2.36.0 للمصادقة والملفات والخدمات الخلفية.",
            "SDWebImageSwiftUI 3.1.4 للصور الشبكية والتخزين المؤقت.",
            "swift-system كاعتماد مباشر/انتقالي ضمن الرسم الحالي.",
            "SpotifyiOS.framework مضمّن كإطار vendor، مع Web API وPKCE.",
            "FoundationModels وHealthKit وFamilyControls وManagedSettings وDeviceActivity وAlarmKit وActivityKit وWidgetKit وCharts وRealityKit من أطر Apple.",
            "Firebase Core/Crashlytics اختياري عبر canImport؛ عدم ربط SDK يحول الخدمة إلى no-op مع بقاء السجل المحلي.",
        ],
        bullet_num,
    )

    heading(doc, "Feature Flags الحالية", 1)
    table(
        doc,
        ["العلم", "القيمة", "الأثر"],
        [
            ["AIQO_CHAT_V1_1_ENABLED", "true", "مسار الدردشة الحديث."],
            ["CAPTAIN_BRAIN_V2_ENABLED", "true", "عقل الكابتن الطبقي."],
            ["GEMINI_3_PREVIEW", "false", "منع preview واستخدام المستقر."],
            ["CAPTAIN_REAL_STREAMING", "false", "عدم فرض streaming الحقيقي."],
            ["CAPTAIN_VOICE_CLOUD", "true", "إتاحة مسار الصوت السحابي."],
            ["CRISIS_DETECTOR", "true", "تفعيل كشف الأزمة."],
            ["HIGH_FIDELITY_3D", "true", "مسار العرض عالي الدقة."],
            ["KERNEL", "true", "إظهار Kernel."],
            ["MEMORY_V4", "true", "مخازن الذاكرة الحديثة/V5."],
            ["NOTIFICATION_BRAIN", "true", "التنبيهات الذكية."],
            ["LEARNING_V2 / STAGE2", "true / true", "Learning Spark الحديث."],
            ["PLANK_LADDER", "false", "إخفاء بديل البلانك."],
            ["ON_DEVICE_VERIFICATION", "true", "OCR/AI محلي للشهادات."],
            ["SMART_WATER", "true", "الماء الذكي."],
            ["TRIBE_VISIBLE / BACKEND", "false / false", "إخفاء Tribe وتعطيل خادمه."],
        ],
        [4050, 1260, 4050],
        font_size=8.2,
    )
    paras(
        doc,
        """
        توجد أعلام إضافية للثقافة والعاطفة والاستباق والذاكرة، وعدة مفاتيح لاستخدام Cloud Proxy. بعض قيم الوكيل تأتي من Secrets.xcconfig أو إعداد Build لا من Info.plist الثابت؛ لذلك لا يمكن من المستودع وحده الجزم بأن كل استدعاء إنتاج يمر بالخادم أو أن المسار المباشر مغلق. يجب توثيق مصفوفة Build لكل بيئة.
        """,
    )

    heading(doc, "الخلفية والروابط", 2)
    bullets(
        doc,
        [
            "Background modes: audio، remote notifications، fetch، processing، location.",
            "BG task identifiers تشمل تحديث الإشعارات، الخمول، وتجميع Brain الليلي.",
            "URL schemes: aiqo وaiqo-spotify.",
            "Entitlements: HealthKit، Sign in with Apple، Siri، Family Controls، APS production، وApp Groups.",
            "App Groups الأساسية: group.com.aiqo.kernel2 وgroup.aiqo.",
        ],
        bullet_num,
    )


def _add_data_security_subscriptions_base(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "التاسع",
        "البيانات، الخصوصية، الأمان، والاشتراكات",
        "من يجمع ماذا، أين تتحرك البيانات، كيف تُفتح الميزات، وأين تقع المخاطر التشغيلية.",
    )
    heading(doc, "تصنيف البيانات", 1)
    table(
        doc,
        ["الفئة", "أمثلة", "المعالجة الأساسية"],
        [
            ["صحة ولياقة", "نوم، قلب، خطوات، ماء، تمارين.", "HealthKit ومحلي؛ ملخصات منتقاة للذكاء."],
            ["هوية", "اسم، بريد Apple/Supabase، User ID.", "Supabase وKeychain/حالة الحساب."],
            ["محتوى مستخدم", "محادثة، ملاحظات، تذكيرات.", "محلي وCloud بعد الموافقة للمدفوع."],
            ["صور", "ثلاجة، جسم، تقدم، شهادة.", "محلي افتراضيًا؛ مسارات محددة فقط للسحابة."],
            ["سلوك", "أحداث واجهة، تقدم، اشتراك.", "JSONL محلي؛ مزود خارجي قابل للحقن."],
            ["أعطال", "stack، سياق منقح، جهاز/نسخة.", "محلي وCrashlytics إن كان مربوطًا."],
            ["انتباه رقمي", "tokens لاختيارات FamilyControls.", "App Group محلي، لا رفع نصي."],
        ],
        [1870, 3120, 4370],
        font_size=8.5,
    )
    paras(
        doc,
        """
        PrivacyInfo.xcprivacy يعلن Fitness وHealth وUser Content وName وEmail وUser ID وPhotos/Videos كبيانات مرتبطة بالمستخدم، مع عدم استخدام للتتبع. أسباب API المعلنة تشمل UserDefaults CA92.1 وFile Timestamp 0A2A.1. هذا البيان يجب أن يظل متزامنًا مع أي SDK أو مزود تحليلات جديد.
        """,
    )


def add_web_public_surface(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "العاشر",
        "الموقع والسطح العام للذكاء",
        "ما تقدمه منظومة aiqo-web للعامة، وكيف تختلف عن بيانات التطبيق الخاصة.",
    )
    heading(doc, "aiqo-web", 1)
    paras(
        doc,
        """
        الموقع مبني بـNext.js 16.2.3 وReact 19.2.4 وTailwind 4. هو واجهة تسويق ودعم وقانوني، لكنه يحتوي أيضًا سطحًا عامًا للآلات: ملفات معرفة JSON، بحث، llms.txt، ai-plugin، OpenAPI، robots، وخادم MCP يعمل عبر stdio.

        هذا السطح يشرح قدرات AiQo ولا يصل إلى HealthKit أو ذاكرة المستخدم. البحث العام keyword-based داخل محتوى محلي، لا LLM ولا حساب مستخدم. هذه الحدود جيدة: قابلية اكتشاف المنتج لا ينبغي أن تتحول إلى بوابة بيانات صحية.
        """,
    )
    table(
        doc,
        ["المسار/السطح", "الغرض", "حساسية البيانات"],
        [
            ["/ai/*.json", "معرفة منظمة عن المنتج.", "عامة"],
            ["/api/knowledge/search", "بحث كلمات في المعرفة العامة.", "عامة؛ بلا حساب"],
            ["llms.txt", "فهرس إرشادي لوكلاء الذكاء.", "عامة"],
            ["ai-plugin + OpenAPI", "وصف تكامل قابل للقراءة.", "عامة"],
            ["MCP stdio", "7 أدوات وResource للمعرفة.", "محلي/تشغيلي؛ لا بيانات مستخدم"],
            ["privacy/terms/support", "السياسات والمساعدة.", "عامة"],
        ],
        [2800, 4060, 2500],
    )
    heading(doc, "دور الموقع في الحقيقة المنتجية", 2)
    bullets(
        doc,
        [
            "الموقع قناة وعد، لكنه ليس مصدرًا تلقائيًا للـ entitlements أو الأرقام.",
            "الأسعار يجب أن تأتي من StoreKit أو رابط المتجر حسب البلد، لا نص ثابت.",
            "عبارات الخصوصية يجب أن تعكس مسارات proxy/direct الفعلية وموافقة السحابة.",
            "مدد Peaks وحدود الذاكرة والتجربة تحتاج مزامنة مع الكود قبل كل نشر.",
            "يمكن توليد صفحات المقارنة وملفات المعرفة من EntitlementPolicy موحدة لتقليل الانجراف.",
        ],
        bullet_num,
    )
    callout(
        doc,
        "فصل حاسم",
        "Public Knowledge يشرح AiQo؛ Captain Memory تعرف المستخدم. يجب ألا يختلط المساران في API أو logging أو caching.",
        "mint",
    )


def add_quality_and_risks(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "الحادي عشر",
        "الاختبارات، الرصد، الأداء، والمخاطر",
        "تقييم قابلية الوثوق: ما الذي تختبره المنظومة، أين توجد نقاط قوة، وما الذي يحتاج إغلاقًا قبل التوسع.",
    )
    heading(doc, "الاختبارات وCI", 1)
    paras(
        doc,
        """
        يحتوي المستودع على نحو 69 ملف اختبار Swift وقرابة 78 XCTestCase/مجموعة و542 دالة test أو @Test بحسب العد النصي. التغطية موضوعية أكثر من كونها سطحية: الكابتن، الخصوصية، الذاكرة، الماء، المشتريات، QuestKit، الصوت، النوم، Kernel، والـ stores. توجد Targets لاختبارات iPhone وWatch وUI.

        GitHub Actions على macos-26 ينشئ Secrets مؤقتة، يختار Xcode 26، ويشغل build-for-testing كبوابة صلبة. تشغيل الاختبارات نفسه continue-on-error بسبب تقلبات معروفة في مسارات الصوت/المحاكاة؛ لذلك نجاح CI لا يعني أن كل test مر. يجب عرض عدد الفشل كإشارة، ثم جعل suites المستقرة حاجبة تدريجيًا.
        """,
    )
    heading(doc, "الرصد والتحليلات", 2)
    bullets(
        doc,
        [
            "AnalyticsService يثري الحدث بنوع الجهاز والنظام والنسخة واللغة والمنطقة الزمنية.",
            "في Debug يطبع Console provider، وفي كل بناء يحفظ LocalAnalyticsProvider حتى 5000 سجل JSONL.",
            "مزود تحليلات خارجي يمكن تسجيله عبر البروتوكول، لكن لا يوجد مزود افتراضي صريح في الملف المفحوص.",
            "CrashReporter يحفظ حتى 50 crash/non-fatal محليًا ويجسرها إلى CrashReportingService.",
            "Crashlytics مشروط بـcanImport، ويعقم الأخطاء قبل الإرسال؛ عند غياب SDK يبقى المسار المحلي.",
        ],
        bullet_num,
    )

    heading(doc, "الأداء والاستدامة", 2)
    paras(
        doc,
        """
        توجد عدة آليات واعية بالطاقة: تهدئة Watch snapshots والويدجت وHealthKit observers، ميزانيات تنبيه، ضغط محادثة حتمي، Actor stores، وتخزين مؤقت للصوت والصور. في المقابل، بعض Stores القديمة تستخدم DispatchQueue.sync على MainActor، وبعض عمليات SwiftData الثقيلة مؤجلة يدويًا بعد Task.yield كي ترسم فقاعة الكتابة.

        لا توجد في الأدلة المفحوصة نتائج Instruments حديثة تربط زمن الإطلاق والذاكرة والطاقة بأجهزة حقيقية. تقارير الأداء النصية مفيدة للتصميم لكنها ليست قياسًا. قبل إطلاق واسع، يلزم baseline على جهاز Apple Intelligence مؤهل وآخر غير مؤهل، وساعة، واتصال ضعيف.
        """,
    )

    heading(doc, "أبرز الديون التقنية", 1)
    table(
        doc,
        ["المجال", "الدين", "الأثر"],
        [
            ["الاشتراك", "TierGate + AccessManager + SubscriptionTier.", "وصول وأسعار وحدود متضاربة."],
            ["HealthKit", "HealthKitService + HealthKitManager.", "صلاحيات وتجميع غير موحد."],
            ["المستويات", "LevelStore + LevelSystem.", "رتب وشارات مختلفة."],
            ["التحديات", "QuestKit + نماذج Challenge أقدم.", "مسارات orphan أو حساب مزدوج."],
            ["الذاكرة", "ترحيلات V3/V5 وshadow-write.", "فقد صامت أو صعوبة تشخيص."],
            ["Foundation Models", "Helpers متقاربة في وحدات عدة.", "سلوك fallback غير موحد."],
            ["الإشعارات", "قيم ميزانية وهدوء متعددة.", "وعد غير متسق وإزعاج محتمل."],
            ["الوثائق", "موقع/Blueprint/كود بإيقاعات مختلفة.", "تسويق ودعم غير دقيق."],
        ],
        [1870, 3740, 3750],
        font_size=8.3,
    )

    heading(doc, "مخاطر مرتبة بالأولوية", 1)
    table(
        doc,
        ["الأولوية", "الخطر", "المعالجة المقترحة"],
        [
            ["P0", "لا توجد بوابة tier ومعدل طلب server-side مؤكدة لوكلاء AI.", "Entitlement lookup + quota + rate limit بالهوية."],
            ["P0", "معلومات سلامة/نوم قد تُوصف بخصوصية أوسع من الواقع.", "لغة دقيقة واختبار consent لكل fallback."],
            ["P1", "تضارب الاشتراك يفتح/يغلق المسار خطأ.", "EntitlementPolicy واحدة واختبارات matrix."],
            ["P1", "Receipt endpoint غير موجود في المستودع.", "توثيق نشره أو إضافته ومراقبة الفشل."],
            ["P1", "MainTab يفرض RTL للإنجليزية.", "اختبار لغة واتجاه على مستوى root."],
            ["P1", "صور التقدم قد تدخل backup دون بيان.", "File protection وbackup policy واضحة."],
            ["P2", "CI tests advisory.", "عزل flaky suites وجعل الباقي hard gate."],
            ["P2", "Tribe backend جاهز جزئيًا ومخفي.", "قرار إطلاق/إزالة مع threat model اجتماعي."],
            ["P2", "مدد وخيارات UI أكبر من entitlement.", "توليد الاختيارات من السياسة نفسها."],
        ],
        [930, 4210, 4220],
        font_size=8.1,
    )

    heading(doc, "خطة إغلاق عملية", 2)
    bullets(
        doc,
        [
            "أسبوع 1: توحيد EntitlementPolicy وكتابة اختبارات Free/Max/Pro/Trial لكل مدخل.",
            "أسبوع 2: tier + quota على Edge Functions، وتأكيد validate-receipt أو حذفه.",
            "أسبوع 3: تدقيق privacy copy ومسارات الصور والنوم والصوت وBuild flags.",
            "أسبوع 4: توحيد HealthKit والمستويات والإشعارات أو وضع طبقة توافق معلنة.",
        ],
        bullet_num,
    )


def add_truth_audit(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "الثاني عشر",
        "تدقيق الحقيقة والفروقات",
        "أين تختلف المواد القديمة أو التسويقية عن الكود الحالي، وما العبارة الدقيقة التي ينبغي اعتمادها.",
    )
    heading(doc, "سجل الفروقات", 1)
    table(
        doc,
        ["الموضوع", "وصف قديم/عام", "الحقيقة التنفيذية بتاريخ الوثيقة"],
        [
            ["التبويبات", "3 تبويبات وكابتن عائم.", "4 تبويبات: الرئيسية، النادي، المطبخ، الكابتن."],
            ["الكابتن المجاني", "الكابتن مدفوع فقط.", "محادثة Apple Intelligence محلية مجانية؛ العمق الدائم مدفوع."],
            ["الذاكرة", "200/500 حقيقة.", "100/500/1000 عرضًا؛ 120/600/1200 سياسة؛ الاسترجاع 0/18/40."],
            ["Kernel", "Max فقط.", "Free لاختيار واحد؛ المدفوع غير محدود."],
            ["التجربة", "7 أيام بلا بطاقة.", "عرض StoreKit أسبوعي ويتطلب عادة وسيلة دفع Apple."],
            ["Peaks", "برامج 4–16 أسبوعًا.", "Seeds الحالية 10–24 أسبوعًا."],
            ["التصدير", "CSV + JSON + PDF لكل الصحة.", "المتحقق تقرير أسبوعي CSV/PDF؛ لا JSON شامل مثبت."],
            ["خصوصية النوم", "لا شيء يغادر الجهاز.", "العينات الخام محلية؛ ملخص مجمع قد يذهب للسحابة بعد موافقة fallback."],
            ["مفاتيح AI", "الخادم فقط دائمًا.", "هناك proxy/direct؛ القيمة الفعلية تتبع إعداد Build."],
            ["التحقق من الإيصال", "خدمة خلفية مكتملة.", "العميل يستدعي endpoint غير موجود داخل المستودع."],
            ["Prompt", "7 طبقات.", "المؤلف الحالي يجمع قرابة 15 كتلة وظيفية."],
            ["Tribe", "ميزة اجتماعية في المنظومة.", "منفذة جزئيًا لكن مخفية والخلفية مغلقة افتراضيًا."],
        ],
        [1560, 3120, 4680],
        font_size=7.9,
    )
    heading(doc, "صياغات دقيقة مقترحة", 2)
    bullets(
        doc,
        [
            "«تحدث مع الكابتن مجانًا على الأجهزة المدعومة؛ افتح الذاكرة والخطط والتخصيص مع Max وPro.»",
            "«بيانات مراحل نومك الخام تبقى على الجهاز. عند تفعيل الذكاء السحابي قد نرسل ملخصًا منقحًا لتحسين التحليل.»",
            "«جرّب الاشتراك أسبوعًا وفق شروط App Store، ثم يتجدد ما لم تُلغِ.»",
            "«Kernel مجاني لاختيار واحد؛ الخطط المدفوعة تفتح اختيارات غير محدودة.»",
            "«Peaks يقدم برامج مرحلية حسب الهدف، وقد تمتد الأهداف الحالية من 10 إلى 24 أسبوعًا.»",
            "«تصدير التقرير الأسبوعي متاح PDF وCSV.»",
        ],
        bullet_num,
    )

    heading(doc, "حدود هذه اللقطة", 2)
    paras(
        doc,
        """
        الفحص يثبت ما في المستودع وإعداداته الظاهرة، لا أسرار الإنتاج ولا نشر Supabase الفعلي ولا حالة Remote Config بعد الإطلاق. كذلك لا يثبت أن النسخة 1.0.8 قُبلت من المتجر؛ الموجود رسميًا وقت التحقق كان 1.0.7، بينما المشروع يحمل 1.0.8 Build 31.

        أي تغيير بعد 3 يوليو 2026—سعر، flag، نموذج، سياسة متجر، endpoint، أو محتوى—قد يجعل جزءًا من الوثيقة تاريخيًا. يُنصح بإعادة توليد «لقطة الحقيقة» مع كل إصدار متجر، لا مع كل commit.
        """,
    )


def add_appendices(doc: Document, bullet_num: int) -> None:
    part(
        doc,
        "—",
        "الملاحق",
        "مرجع سريع للمكونات والبيانات والمصطلحات والمصادر المستخدمة في التحقق، دون الحاجة إلى قراءة السرد الكامل.",
    )
    heading(doc, "ملحق أ — خريطة الوحدات", 1)
    table(
        doc,
        ["الوحدة", "الحجم التقريبي/الملاحظة"],
        [
            ["Features/Captain", "155 ملفًا؛ أكبر وحدة ذكاء وذاكرة."],
            ["Features/Gym", "102 ملف؛ جلسات وخطط وتحديات."],
            ["Features/Tribe", "58 ملفًا؛ مخفي حاليًا."],
            ["Features/Kitchen", "32 ملفًا؛ تغذية وثلاجة وخطة."],
            ["Features/Home", "22 ملفًا؛ لوحة اليوم."],
            ["LegendaryChallenges", "16 ملفًا؛ Peaks."],
            ["Kernel", "16 ملفًا + 3 extension targets."],
            ["DesignSystem", "13 ملفًا؛ tokens ومكونات."],
            ["Notifications", "12 ملفًا؛ brain وجدولة."],
            ["Sleep", "11 ملفًا؛ تحليل ومنبه وAI."],
            ["Onboarding", "8 ملفات رئيسية؛ تدفق متعدد المراحل."],
            ["SmartWater", "7 ملفات؛ هدف ووتيرة وويدجت."],
        ],
        [3740, 5620],
    )

    heading(doc, "ملحق ب — أنواع HealthKit البارزة", 1)
    table(
        doc,
        ["المجموعة", "الأنواع"],
        [
            ["الحركة", "Steps، active energy، walking/running distance، cycling distance، stand hour."],
            ["القلب", "Heart rate، resting HR، HRV، walking HR average، VO₂ max، oxygen saturation."],
            ["النوم", "Sleep analysis ومراحله."],
            ["الجسم", "Body mass، height من الملف/المصدر المناسب."],
            ["الغذاء", "Dietary water."],
            ["التمرين", "HKWorkout وActivity Summary."],
        ],
        [2340, 7020],
    )

    heading(doc, "ملحق ج — مصطلحات", 1)
    table(
        doc,
        ["المصطلح", "المعنى في AiQo"],
        [
            ["Captain", "الشخصية الحوارية وواجهة العقل والذاكرة."],
            ["Brain V2", "طبقات الاستشعار والذاكرة والاستدلال والخصوصية والاستباق."],
            ["Directive", "تعليم مستمر من المستخدم مرتبط بحدث وفعل."],
            ["Bio State", "ملخص إشارات اليوم الصحية والسلوكية."],
            ["Daily Aura", "تمثيل بصري لحالة اليوم وهدفه."],
            ["Kernel", "حجب مشتتات مقابل تحدٍ حركي."],
            ["QuestKit", "محرك المراحل والمهام والترقيات."],
            ["Peak", "برنامج للوصول إلى رقم أسطوري محدد."],
            ["My Vibe", "حالة زمنية وموسيقى محلية/Spotify."],
            ["TierGate", "سياسة وصول أحدث للكابتن والميزات."],
            ["AccessManager", "بوابة أقدم ما زالت مستخدمة في شاشات."],
            ["CloudSafeProfile", "جزء من الملف مسموح للسحابة بعد الموافقة."],
        ],
        [2800, 6560],
    )

    heading(doc, "ملحق د — مصادر التحقق", 1)
    bullets(
        doc,
        [
            "الكود وإعدادات Xcode داخل مستودع /Users/mohammedraad/Desktop/AiQo.",
            "AiQo_Master_Blueprint.md وARCHITECTURE.md وCHANGELOG.md وSECURITY.md مع اعتبارها مصادر مساندة لا نهائية.",
            "Info.plist وentitlements وPrivacyInfo.xcprivacy وStoreKit configuration.",
            "Supabase Edge Functions: captain-chat وcaptain-voice.",
            "الموقع الرسمي: https://aiqo.app/.",
            "صفحة متجر الإمارات: https://apps.apple.com/ae/app/aiqo/id6755132504.",
            "حالة الفحص: 3 يوليو 2026؛ branch program/world-class-completion.",
        ],
        bullet_num,
    )

    heading(doc, "ملحق هـ — قائمة مراجعة قبل كل إصدار", 1)
    bullets(
        doc,
        [
            "طابق MARKETING_VERSION وBuild مع المتجر وCHANGELOG.",
            "اختبر Free/Max/Pro/Trial لكل تبويب وDeep Link وWidget.",
            "راجع كل Feature Flag وBuild flag وبيئة secrets.",
            "اختبر سحب موافقات AI والصوت والصور وHealthKit.",
            "تحقق من tier/rate limits في Edge Functions ومن receipt endpoint.",
            "شغل اختبارات UI للشراء والاستعادة وحذف الحساب.",
            "راجع Privacy Manifest وApp Privacy مع SDKs الفعلية.",
            "اختبر العربية والإنجليزية وRTL/LTR وDynamic Type وVoiceOver.",
            "اختبر جهازًا يدعم Apple Intelligence وآخر لا يدعمه.",
            "حدّث الموقع وملفات AI العامة من مصدر السياسة نفسه.",
        ],
        bullet_num,
    )

    heading(doc, "الخاتمة", 1)
    paras(
        doc,
        """
        AiQo ليس تطبيقًا صغيرًا أضيفت إليه دردشة؛ هو منظومة واسعة تجمع الصحة والسلوك والذكاء والساعة والويب. أقوى ما فيه أن القطع موجودة فعلًا: كابتن محلي وسحابي، ذاكرة عميقة، نادي غني، مطبخ، نوم، ماء، Kernel، Watch، وتقارير. التحدي القادم ليس إضافة مزايا أكثر بقدر ما هو توحيد الحقيقة: بوابة اشتراك واحدة، سياسة خصوصية واحدة، مصادر بيانات محددة، ووعود تسويقية تتولد من الكود.

        إذا تحقق هذا التوحيد، تصبح المنظومة أسهل في الشرح والاختبار والدعم، ويشعر المستخدم أن AiQo «يعرف يومه» دون أن يتجاوز حدوده. ذلك هو جوهر المنتج: ذكاء حاضر بما يكفي ليكون مفيدًا، ومقيد بما يكفي ليبقى موثوقًا.
        """,
    )
    callout(
        doc,
        "الخلاصة في سطر",
        "AiQo هو رفيق صحة عربي متعدد الأسطح؛ قيمته الحقيقية في تحويل البيانات إلى فعل مستمر، ونجاحه طويل الأمد يعتمد على توحيد السياسات بقدر اعتماده على جودة الذكاء.",
        "mint",
    )
    return

    heading(doc, "الموافقات المنفصلة", 1)
    bullets(
        doc,
        [
            "HealthKit: لكل نوع حسب نظام Apple، مع إمكانية الرفض أو السحب.",
            "AI Cloud: قبل نقل النص/الملخصات لمسار الذكاء السحابي.",
            "Voice Cloud: منفصلة عن النص لأنها ترسل نصًا/صوتًا إلى مزود آخر.",
            "Body Photo: قبل استخدام صورة الجسم في الخطة.",
            "Kitchen/Gym Image: موافقة سياقية قبل تحليل صورة.",
            "On-device verification: إعداد مستقل لإثبات الشهادات محليًا.",
            "Notifications، Location، FamilyControls، AlarmKit: كل منها تفويض نظامي منفصل.",
        ],
        bullet_num,
    )
    callout(
        doc,
        "ليس E2E",
        "الاتصال المحمي وKeychain وتعقيم البيانات لا يساوي تشفيرًا طرفًا لطرف للذكاء السحابي. الوصف الأدق: نقل مشفر عبر المنصة، تقليل حمولة، وموافقة قبل مشاركة المحدد.",
        "red",
    )

    heading(doc, "وظائف Supabase الطرفية", 1)
    table(
        doc,
        ["الوظيفة", "ما تفعله", "ضوابط حالية"],
        [
            ["captain-chat", "وكيل Gemini عادي وSSE.", "JWT، حد POST 256KB، allowlist للنموذج."],
            ["captain-voice", "وكيل MiniMax للصوت.", "JWT، حد 16KB، allowlist للنموذج."],
            ["validate-receipt", "يستدعيها ReceiptValidator.", "غير موجودة في المستودع الحالي؛ قد تكون نشرًا خارجيًا."],
        ],
        [2340, 3900, 3120],
    )
    paras(
        doc,
        """
        captain-chat يتحقق من JWT عبر getUser ويمرر gemini-2.5-flash أو preview المسموح. captain-voice يفعل الشيء نفسه للمزود الصوتي. التعليقات في auth helper تذكر أن إعادة فحص المستوى على الخادم عمل مستقبلي؛ أي أن الوكيل لا يثبت tier server-side حاليًا. البوابة الرئيسية داخل العميل، وهذا غير كاف وحده لحماية تكلفة API أو منع عميل معدل.

        لا يظهر في الوظيفتين rate limit لكل مستخدم. حدود حجم الطلب جيدة لكنها لا تمنع الإساءة المتكررة. الأولوية الأمنية: tier enforcement مركزي، rate limiting، quotas، idempotency عند اللزوم، وتسجيل منقح.
        """,
    )

    heading(doc, "الأسرار ومسار الشبكة", 2)
    bullets(
        doc,
        [
            "Secrets.xcconfig محجوب عن Git، وInfo.plist يستخدم placeholders.",
            "يوجد دعم لمسار proxy ومسار مباشر في بعض الخدمات؛ اختيار الإنتاج يعتمد على build flags.",
            "Spotify tokens تحفظ في Keychain؛ callback عبر aiqo-spotify.",
            "أخطاء السحابة يجب ألا تسجل prompt أو token كاملًا؛ PrivacySanitizer يعقم الرسائل.",
            "SECURITY.md يوجه البلاغات إلى support@aiqo.app ويستهدف إقرارًا خلال 72 ساعة، دون برنامج مكافآت معلن.",
        ],
        bullet_num,
    )

    heading(doc, "الاشتراكات ومنطق الوصول", 1)
    table(
        doc,
        ["المستوى", "المعنى العملي"],
        [
            ["Free", "الرئيسية والصحة والماء، كابتن محلي، Kernel لاختيار واحد، وحدود ضيقة."],
            ["Max", "ذاكرة واسترجاع متوسط، خطط أسبوعية، مطبخ/ميزات مدفوعة، إشعارات أعمق."],
            ["Pro", "عمق ذاكرة أكبر، تعدد أسابيع، Peaks كامل، صوت/صور/تحليل أعمق حسب البوابة."],
            ["Trial", "اشتراك StoreKit تمهيدي، effectiveAccessTier يعامله كـPro."],
        ],
        [2340, 7020],
    )
    paras(
        doc,
        """
        المنتجات الفعلية: Max بالمعرف com.mraad5000.aiqo.max—الخطأ المطبعي في 5000 متعمد وغير قابل للتغيير بعد النشر—وPro بالمعرف com.mraad500.aiqo.Intelligence.pro. تحتفظ الشفرة بمعرفات legacy لاستعادة مشتركي الإصدارات السابقة. StoreKit configuration يعرّف اشتراكًا شهريًا وتجربة أسبوع واحد، ولا يفعّل Family Sharing.

        في متجر الإمارات وقت التحقق: Max 39.99 د.إ وPro 79.99 د.إ شهريًا. fallback داخل التطبيق 9.99 و19.99 دولار، لكنه ليس مصدر السعر المعروض النهائي؛ StoreKit هو المرجع. عبارة الموقع «تجربة 7 أيام بلا بطاقة» لا تتطابق مع تجربة App Store المعتادة التي تتطلب وسيلة دفع بالحساب.
        """,
    )
    heading(doc, "بوابات متداخلة", 2)
    paras(
        doc,
        """
        TierGate هو المرجع الأحدث لعقل الكابتن: basicLife وbasicNotifications مجانية؛ captainChat والميموري والتوجيهات Max؛ multi-week فوق أسبوع وweekly insights وmonthly reflection وphoto analysis وpremium voice والثقافة Pro. لكنه يعرّف requiredTier(.kernel)=Max في موضع، بينما kernelAppLimit يتيح مجانًا تطبيقًا واحدًا والواجهة الحالية تفتح Kernel للجميع.

        AccessManager أقدم ويضع Captain وGym وKitchen وMy Vibe وChallenges وDataTracking خلف Max، مع أرقام ذاكرة 200/500. بعض الشاشات ما زالت تستخدمه. SubscriptionTier وNotificationBrain يحملان بدورهما أرقامًا أخرى. النتيجة: بوابة المستخدم قد تختلف حسب المسار، وهذا أهم دين منتجي في الإصدار الحالي.
        """,
    )
    table(
        doc,
        ["البعد", "Free", "Max", "Pro/Trial"],
        [
            ["حقائق الذاكرة المعروضة", "100", "500", "1000"],
            ["سقف TierGate الدلالي", "120", "600", "1200"],
            ["عمق الاسترجاع", "0", "18", "40"],
            ["Context tokens", "0 سحابيًا", "8000", "32000"],
            ["أسابيع الخطة", "—/أساسي محلي", "1", "4 حسب TierGate"],
            ["Kernel selections", "1", "غير محدود", "غير محدود"],
            ["Callback memory", "—", "30 يومًا", "غير محدود"],
        ],
        [3120, 1870, 1870, 2500],
        font_size=8.2,
    )
    callout(
        doc,
        "قرار مطلوب",
        "اجعل EntitlementPolicy واحدة قابلة للاختبار، وتُستدعى من الواجهة والخدمة والكابتن والخادم. أي جدول أسعار أو سعة يجب أن يُولد منها لا أن يُنسخ يدويًا.",
        "gold",
    )

    heading(doc, "التجربة القديمة والجديدة", 2)
    paras(
        doc,
        """
        FreeTrialManager القديم لا يزال يحتفظ بمرساة سبعة أيام في Keychain/UserDefaults، لكنه ليس المسار الجديد لبدء التجربة. onboarding الحالي لا يبدأ trial بلا بطاقة؛ PurchaseManager وStoreKit هما المسار. يجب إزالة أو عزل الكود القديم كي لا يخلق entitlement شبحًا أو تحليلات متضاربة.

        PurchaseManager يستخدم StoreKit 2 كمرجع entitlement محلي ويستعيد المشتريات. ReceiptValidator يحاول التحقق غير الحاجب عبر وظيفة Supabase. غياب الوظيفة من المستودع يعني أن التحقق قد يفشل بهدوء أو يعتمد على نشر غير موثق؛ لا ينبغي أن يحجب المستخدم الشرعي، لكن يجب قياسه.
        """,
    )

    heading(doc, "الشخصية والنبرة", 2)
    paras(
        doc,
        """
        «الكابتن حمودي» ليس مجرد اسم لواجهة الدردشة؛ هو طبقة الهوية التي توحّد اللغة والتحفيز. النبرة الأساسية عراقية دافئة، مباشرة، غير متصنعة، وتوازن بين الصرامة والرعاية. في المستويات المدفوعة يمكن اختيار أساليب عملية، حنونة، صارمة، تحليلية، رؤيوية، أو مرشدة، مع أسلوب مخصص في Pro. أما المجاني فيستخدم صوتًا ثابتًا بسيطًا مع اسم المستخدم فقط.

        المنتج عربي أولًا لكن ليس عربيًا فقط. يدعم الإنجليزية على مستوى الواجهة والإشعارات ومخرجات الذكاء. يجب الانتباه إلى أن MainTabScreen يفرض RTL حاليًا حتى عند اختيار الإنجليزية؛ هذه نقطة تنفيذية قد تجعل بعض محاذاة الإنجليزية غير طبيعية رغم وجود نظام توطين كامل.
        """,
    )

    heading(doc, "الحلقة اليومية الأساسية", 2)
    flow(
        doc,
        "قياس HealthKit وWatch ← تلخيص الرئيسية ← إشارة أو تنبيه ذكي ← فعل: تمرين/ماء/نوم/تركيز ← XP وسلسلة ← ذاكرة وتحليل ← خطوة اليوم التالي",
    )
    paras(
        doc,
        """
        نجاح AiQo لا يعتمد على أن يفتح المستخدم كل ميزة. الحلقة المثالية قصيرة: يرى حالته في الرئيسية، يلتقط توجيهًا واحدًا، ينفذ فعلًا، ثم يحصل على أثر واضح. الطبقات الأعمق—الكابتن، الخطط، المقارنات، وPeaks—تزيد القيمة للمستخدم الملتزم، بينما الويدجت وWatch والإشعارات تحافظ على الوجود دون فتح التطبيق.
        """,
    )

    heading(doc, "ما لا ينبغي أن يُفهم خطأ", 2)
    bullets(
        doc,
        [
            "AiQo ليس شبكة اجتماعية عامة الآن؛ Tribe مخفي وإسناده الخلفي مغلق افتراضيًا.",
            "AiQo لا يثبت سببية طبية من الترابطات؛ تقارير النوم والتعافي والتوتر تقديرات إرشادية.",
            "النسخة المجانية لا تملك ذاكرة دائمة للكابتن رغم بقاء سياق المحادثة الحالية على الجهاز.",
            "وجود كود ميزة لا يعني ظهورها؛ الأعلام وإعدادات الاشتراك وBuild Configuration تغير السلوك.",
            "الموقع والمتجر ليسا دائمًا متزامنين مع الفرع الحالي؛ يوجد قسم تدقيق كامل للفروقات في آخر الوثيقة.",
        ],
        bullet_num,
    )


def add_executive_snapshot(doc: Document, bullet_num: int) -> None:
    _add_executive_snapshot_base(doc, bullet_num)

    heading(doc, "الشخصية والنبرة", 2)
    paras(
        doc,
        """
        «الكابتن حمودي» ليس مجرد اسم لواجهة الدردشة؛ هو طبقة الهوية التي توحّد اللغة والتحفيز. النبرة الأساسية عراقية دافئة، مباشرة، غير متصنعة، وتوازن بين الصرامة والرعاية. في المستويات المدفوعة يمكن اختيار أساليب عملية، حنونة، صارمة، تحليلية، رؤيوية، أو مرشدة، مع أسلوب مخصص في Pro. أما المجاني فيستخدم صوتًا ثابتًا بسيطًا مع اسم المستخدم فقط.

        المنتج عربي أولًا لكن ليس عربيًا فقط. يدعم الإنجليزية على مستوى الواجهة والإشعارات ومخرجات الذكاء. يجب الانتباه إلى أن MainTabScreen يفرض RTL حاليًا حتى عند اختيار الإنجليزية؛ هذه نقطة تنفيذية قد تجعل بعض محاذاة الإنجليزية غير طبيعية رغم وجود نظام توطين كامل.
        """,
    )

    heading(doc, "الحلقة اليومية الأساسية", 2)
    flow(
        doc,
        "قياس HealthKit وWatch ← تلخيص الرئيسية ← إشارة أو تنبيه ذكي ← فعل: تمرين/ماء/نوم/تركيز ← XP وسلسلة ← ذاكرة وتحليل ← خطوة اليوم التالي",
    )
    paras(
        doc,
        """
        نجاح AiQo لا يعتمد على أن يفتح المستخدم كل ميزة. الحلقة المثالية قصيرة: يرى حالته في الرئيسية، يلتقط توجيهًا واحدًا، ينفذ فعلًا، ثم يحصل على أثر واضح. الطبقات الأعمق—الكابتن، الخطط، المقارنات، وPeaks—تزيد القيمة للمستخدم الملتزم، بينما الويدجت وWatch والإشعارات تحافظ على الوجود دون فتح التطبيق.
        """,
    )

    heading(doc, "ما لا ينبغي أن يُفهم خطأ", 2)
    bullets(
        doc,
        [
            "AiQo ليس شبكة اجتماعية عامة الآن؛ Tribe مخفي وإسناده الخلفي مغلق افتراضيًا.",
            "AiQo لا يثبت سببية طبية من الترابطات؛ تقارير النوم والتعافي والتوتر تقديرات إرشادية.",
            "النسخة المجانية لا تملك ذاكرة دائمة للكابتن رغم بقاء سياق المحادثة الحالية على الجهاز.",
            "وجود كود ميزة لا يعني ظهورها؛ الأعلام وإعدادات الاشتراك وBuild Configuration تغير السلوك.",
            "الموقع والمتجر ليسا دائمًا متزامنين مع الفرع الحالي؛ يوجد قسم تدقيق كامل للفروقات في آخر الوثيقة.",
        ],
        bullet_num,
    )


def add_data_security_subscriptions(doc: Document, bullet_num: int) -> None:
    _add_data_security_subscriptions_base(doc, bullet_num)

    heading(doc, "الموافقات المنفصلة", 1)
    bullets(
        doc,
        [
            "HealthKit: لكل نوع حسب نظام Apple، مع إمكانية الرفض أو السحب.",
            "AI Cloud: قبل نقل النص/الملخصات لمسار الذكاء السحابي.",
            "Voice Cloud: منفصلة عن النص لأنها ترسل نصًا/صوتًا إلى مزود آخر.",
            "Body Photo: قبل استخدام صورة الجسم في الخطة.",
            "Kitchen/Gym Image: موافقة سياقية قبل تحليل صورة.",
            "On-device verification: إعداد مستقل لإثبات الشهادات محليًا.",
            "Notifications، Location، FamilyControls، AlarmKit: كل منها تفويض نظامي منفصل.",
        ],
        bullet_num,
    )
    callout(
        doc,
        "ليس E2E",
        "الاتصال المحمي وKeychain وتعقيم البيانات لا يساوي تشفيرًا طرفًا لطرف للذكاء السحابي. الوصف الأدق: نقل مشفر عبر المنصة، تقليل حمولة، وموافقة قبل مشاركة المحدد.",
        "red",
    )

    heading(doc, "وظائف Supabase الطرفية", 1)
    table(
        doc,
        ["الوظيفة", "ما تفعله", "ضوابط حالية"],
        [
            ["captain-chat", "وكيل Gemini عادي وSSE.", "JWT، حد POST 256KB، allowlist للنموذج."],
            ["captain-voice", "وكيل MiniMax للصوت.", "JWT، حد 16KB، allowlist للنموذج."],
            ["validate-receipt", "يستدعيها ReceiptValidator.", "غير موجودة في المستودع الحالي؛ قد تكون نشرًا خارجيًا."],
        ],
        [2340, 3900, 3120],
    )
    paras(
        doc,
        """
        captain-chat يتحقق من JWT عبر getUser ويمرر gemini-2.5-flash أو preview المسموح. captain-voice يفعل الشيء نفسه للمزود الصوتي. التعليقات في auth helper تذكر أن إعادة فحص المستوى على الخادم عمل مستقبلي؛ أي أن الوكيل لا يثبت tier server-side حاليًا. البوابة الرئيسية داخل العميل، وهذا غير كاف وحده لحماية تكلفة API أو منع عميل معدل.

        لا يظهر في الوظيفتين rate limit لكل مستخدم. حدود حجم الطلب جيدة لكنها لا تمنع الإساءة المتكررة. الأولوية الأمنية: tier enforcement مركزي، rate limiting، quotas، idempotency عند اللزوم، وتسجيل منقح.
        """,
    )

    heading(doc, "الأسرار ومسار الشبكة", 2)
    bullets(
        doc,
        [
            "Secrets.xcconfig محجوب عن Git، وInfo.plist يستخدم placeholders.",
            "يوجد دعم لمسار proxy ومسار مباشر في بعض الخدمات؛ اختيار الإنتاج يعتمد على build flags.",
            "Spotify tokens تحفظ في Keychain؛ callback عبر aiqo-spotify.",
            "أخطاء السحابة يجب ألا تسجل prompt أو token كاملًا؛ PrivacySanitizer يعقم الرسائل.",
            "SECURITY.md يوجه البلاغات إلى support@aiqo.app ويستهدف إقرارًا خلال 72 ساعة، دون برنامج مكافآت معلن.",
        ],
        bullet_num,
    )

    heading(doc, "الاشتراكات ومنطق الوصول", 1)
    table(
        doc,
        ["المستوى", "المعنى العملي"],
        [
            ["Free", "الرئيسية والصحة والماء، كابتن محلي، Kernel لاختيار واحد، وحدود ضيقة."],
            ["Max", "ذاكرة واسترجاع متوسط، خطط أسبوعية، مطبخ/ميزات مدفوعة، إشعارات أعمق."],
            ["Pro", "عمق ذاكرة أكبر، تعدد أسابيع، Peaks كامل، صوت/صور/تحليل أعمق حسب البوابة."],
            ["Trial", "اشتراك StoreKit تمهيدي، effectiveAccessTier يعامله كـPro."],
        ],
        [2340, 7020],
    )
    paras(
        doc,
        """
        المنتجات الفعلية: Max بالمعرف com.mraad5000.aiqo.max—الخطأ المطبعي في 5000 متعمد وغير قابل للتغيير بعد النشر—وPro بالمعرف com.mraad500.aiqo.Intelligence.pro. تحتفظ الشفرة بمعرفات legacy لاستعادة مشتركي الإصدارات السابقة. StoreKit configuration يعرّف اشتراكًا شهريًا وتجربة أسبوع واحد، ولا يفعّل Family Sharing.

        في متجر الإمارات وقت التحقق: Max 39.99 د.إ وPro 79.99 د.إ شهريًا. fallback داخل التطبيق 9.99 و19.99 دولار، لكنه ليس مصدر السعر المعروض النهائي؛ StoreKit هو المرجع. عبارة الموقع «تجربة 7 أيام بلا بطاقة» لا تتطابق مع تجربة App Store المعتادة التي تتطلب وسيلة دفع بالحساب.
        """,
    )
    heading(doc, "بوابات متداخلة", 2)
    paras(
        doc,
        """
        TierGate هو المرجع الأحدث لعقل الكابتن: basicLife وbasicNotifications مجانية؛ captainChat والميموري والتوجيهات Max؛ multi-week فوق أسبوع وweekly insights وmonthly reflection وphoto analysis وpremium voice والثقافة Pro. لكنه يعرّف requiredTier(.kernel)=Max في موضع، بينما kernelAppLimit يتيح مجانًا تطبيقًا واحدًا والواجهة الحالية تفتح Kernel للجميع.

        AccessManager أقدم ويضع Captain وGym وKitchen وMy Vibe وChallenges وDataTracking خلف Max، مع أرقام ذاكرة 200/500. بعض الشاشات ما زالت تستخدمه. SubscriptionTier وNotificationBrain يحملان بدورهما أرقامًا أخرى. النتيجة: بوابة المستخدم قد تختلف حسب المسار، وهذا أهم دين منتجي في الإصدار الحالي.
        """,
    )
    table(
        doc,
        ["البعد", "Free", "Max", "Pro/Trial"],
        [
            ["حقائق الذاكرة المعروضة", "100", "500", "1000"],
            ["سقف TierGate الدلالي", "120", "600", "1200"],
            ["عمق الاسترجاع", "0", "18", "40"],
            ["Context tokens", "0 سحابيًا", "8000", "32000"],
            ["أسابيع الخطة", "—/أساسي محلي", "1", "4 حسب TierGate"],
            ["Kernel selections", "1", "غير محدود", "غير محدود"],
            ["Callback memory", "—", "30 يومًا", "غير محدود"],
        ],
        [3120, 1870, 1870, 2500],
        font_size=8.2,
    )
    callout(
        doc,
        "قرار مطلوب",
        "اجعل EntitlementPolicy واحدة قابلة للاختبار، وتُستدعى من الواجهة والخدمة والكابتن والخادم. أي جدول أسعار أو سعة يجب أن يُولد منها لا أن يُنسخ يدويًا.",
        "gold",
    )

    heading(doc, "التجربة القديمة والجديدة", 2)
    paras(
        doc,
        """
        FreeTrialManager القديم لا يزال يحتفظ بمرساة سبعة أيام في Keychain/UserDefaults، لكنه ليس المسار الجديد لبدء التجربة. onboarding الحالي لا يبدأ trial بلا بطاقة؛ PurchaseManager وStoreKit هما المسار. يجب إزالة أو عزل الكود القديم كي لا يخلق entitlement شبحًا أو تحليلات متضاربة.

        PurchaseManager يستخدم StoreKit 2 كمرجع entitlement محلي ويستعيد المشتريات. ReceiptValidator يحاول التحقق غير الحاجب عبر وظيفة Supabase. غياب الوظيفة من المستودع يعني أن التحقق قد يفشل بهدوء أو يعتمد على نشر غير موثق؛ لا ينبغي أن يحجب المستخدم الشرعي، لكن يجب قياسه.
        """,
    )


def build_document() -> None:
    doc, bullet_num, decimal_num = setup_document()
    add_cover(doc)
    add_front_matter(doc, bullet_num)
    add_executive_snapshot(doc, bullet_num)
    add_user_journey(doc, bullet_num, decimal_num)
    add_home_health_sleep(doc, bullet_num)
    add_gym(doc, bullet_num, decimal_num)
    add_kitchen_vibe_kernel(doc, bullet_num)
    add_captain(doc, bullet_num, decimal_num)
    add_watch_widgets_reports(doc, bullet_num)
    add_architecture(doc, bullet_num, decimal_num)
    add_data_security_subscriptions(doc, bullet_num)
    add_web_public_surface(doc, bullet_num)
    add_quality_and_risks(doc, bullet_num)
    add_truth_audit(doc, bullet_num)
    add_appendices(doc, bullet_num)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
